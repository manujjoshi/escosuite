import streamlit as st
import os
import json
import torch
import pandas as pd
from doctr.models import ocr_predictor
from doctr.io import DocumentFile
from dotenv import load_dotenv
from docx import Document as DocxDocument
import fitz  # PyMuPDF for PDF text extraction
import re
from typing import List, Dict

# ----------------- OCR -------------------
@st.cache_resource
def load_doctr_model(device: str = "cuda"):
    try:
        return ocr_predictor(pretrained=True).to(device)
    except Exception:
        return ocr_predictor(pretrained=True).to("cpu")

def extract_pages_with_doctr(pdf_bytes: bytes, model) -> List[str]:
    """Return a list of page-level texts using docTR OCR."""
    doc = DocumentFile.from_pdf(pdf_bytes)
    result = model(doc)
    return [page.render() for page in result.pages]

def extract_page_texts(file, doctr_model) -> List[str]:
    """
    Try to extract selectable text with PyMuPDF per-page.
    If no text is found (scanned PDF), OCR page-by-page with docTR.
    """
    pages = []
    try:
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                t = page.get_text("text")
                pages.append(t or "")
    except Exception:
        pages = []
    # If most pages are empty => OCR fallback
    if not pages or sum(1 for p in pages if p.strip()) < max(1, int(0.2 * len(pages))):
        file.seek(0)
        pdf_bytes = file.read()
        pages = extract_pages_with_doctr(pdf_bytes, doctr_model)
    file.seek(0)
    return pages

# ----------------- Simple Text Analysis (NO API) -------------------
def extract_basic_info(text: str, filename: str) -> dict:
    """
    Extract basic information from RFP text using keyword matching.
    NO API calls - pure text analysis.
    """
    text_lower = text.lower()
    
    # Extract dates (simple regex)
    date_patterns = [
        r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',  # MM/DD/YYYY or DD-MM-YYYY
        r'\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2},?\s+\d{4}\b'
    ]
    dates = []
    for pattern in date_patterns:
        dates.extend(re.findall(pattern, text, re.IGNORECASE))
    
    # Extract potential costs/budgets
    budget_pattern = r'\$[\d,]+(?:\.\d{2})?'
    budgets = re.findall(budget_pattern, text)
    
    # Detect project type keywords
    project_types = []
    type_keywords = {
        'lighting': ['lighting', 'led', 'illumination', 'fixture', 'luminaire', 'lamp'],
        'hvac': ['hvac', 'heating', 'ventilation', 'air conditioning', 'cooling'],
        'solar': ['solar', 'photovoltaic', 'pv system'],
        'water': ['water', 'plumbing', 'water management'],
        'wastewater': ['wastewater', 'waste water', 'sewage'],
        'building envelope': ['building envelope', 'roofing', 'insulation', 'windows'],
        'esco': ['esco', 'energy service', 'performance contract'],
        'generator': ['generator', 'emergency power', 'backup power']
    }
    
    for proj_type, keywords in type_keywords.items():
        if any(kw in text_lower or kw in filename.lower() for kw in keywords):
            project_types.append(proj_type.upper())
    
    # Detect scope keywords
    has_supply = any(kw in text_lower for kw in ['supply', 'furnish', 'provide', 'material'])
    has_installation = any(kw in text_lower for kw in ['install', 'installation', 'labor'])
    has_substitution = any(kw in text_lower for kw in ['substitution allowed', 'or equal', 'approved equal', 'or equivalent'])
    has_no_substitution = any(kw in text_lower for kw in ['no substitution', 'as specified', 'exact match'])
    
    return {
        'project_type': ', '.join(project_types) if project_types else 'General/Unknown',
        'dates_found': dates[:10],  # First 10 dates
        'budgets_found': budgets[:10],  # First 10 budget mentions
        'scope_supply': has_supply,
        'scope_installation': has_installation,
        'scope_substitution_allowed': has_substitution,
        'scope_no_substitution': has_no_substitution,
        'total_pages': 0,  # Will be set later
        'total_chars': len(text),
        'filename': filename
    }

# ----------------- Criteria 1: Scope and Type Pre-Weighting (10%) -------------------
def apply_scope_preweights(rfp_info: dict, full_text: str) -> dict:
    """
    Apply Criteria 1 pre-weighting (10%) dynamically based on detected project type
    and scope. NO API - pure keyword matching.
    """
    LIGHTING_KEYWORDS = ["lighting", "led", "illumination", "fixture", "luminaire"]
    PROJECT_TYPE_KEYWORDS = {
        "hvac": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "solar": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "water": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "waste water": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "wastewater": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "building envelope": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "esco": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "energy saving": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "generator": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
    }

    # Combine filename + text for detection
    filename_part = rfp_info.get('filename', '').lower()
    text = (filename_part + " " + full_text).lower()

    preweights = {"IKIO": 0, "METCO": 0, "SUNSPRINT": 0}
    detected_type, condition = "Unknown", "N/A"
    detection_source = "N/A"

    # --- Lighting branch ---
    if any(k in text for k in LIGHTING_KEYWORDS):
        detected_type = "Lighting"
        has_supply = rfp_info.get('scope_supply', False)
        has_install = rfp_info.get('scope_installation', False)
        substitution_allowed = rfp_info.get('scope_substitution_allowed', False)
        no_substitution = rfp_info.get('scope_no_substitution', False)

        # Determine detection source
        if any(kw in filename_part for kw in LIGHTING_KEYWORDS):
            detection_source = "Filename + Content"
        else:
            detection_source = "Content"

        if has_supply and not has_install and substitution_allowed:
            preweights = {"IKIO": 10, "METCO": 0, "SUNSPRINT": 0}
            condition = "Supply + Substitution Allowed"
        elif has_supply and has_install and substitution_allowed:
            preweights = {"IKIO": 10, "METCO": 10, "SUNSPRINT": 10}
            condition = "Supply + Installation + Substitution Allowed"
        elif has_supply and has_install and no_substitution:
            preweights = {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10}
            condition = "Supply + Installation + Substitution Not Allowed"
        elif not has_supply and has_install:
            preweights = {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10}
            condition = "Installation Only"
        else:
            preweights = {"IKIO": 10, "METCO": 10, "SUNSPRINT": 10}
            condition = "Lighting (Default All)"
    else:
        # --- Non-lighting dynamic detection ---
        for kw, weights in PROJECT_TYPE_KEYWORDS.items():
            if kw in text:
                detected_type = kw.title()
                preweights = weights
                condition = f"Detected via keyword '{kw}'"
                
                # Determine detection source
                if kw in filename_part:
                    detection_source = "Filename + Content"
                else:
                    detection_source = "Content"
                break

    return {
        "detected_type": detected_type,
        "condition": condition,
        "preweights": preweights,
        "detection_source": detection_source
    }

# ----------------- DOCX Export -------------------
def export_to_docx(results: dict, rfp_info: dict, preweight_result: dict, file_path: str):
    doc = DocxDocument()
    doc.add_heading("RFP Analysis Report - Criteria 1 Pre-Weighting", level=1)
    
    # RFP Basic Info
    doc.add_heading("RFP Information", level=2)
    doc.add_paragraph(f"Filename: {rfp_info.get('filename', 'N/A')}")
    doc.add_paragraph(f"Total Pages: {rfp_info.get('total_pages', 0)}")
    doc.add_paragraph(f"Total Characters: {rfp_info.get('total_chars', 0):,}")
    doc.add_paragraph(f"Detected Project Type: {rfp_info.get('project_type', 'N/A')}")
    
    # Scope Detection
    doc.add_heading("Scope Detection", level=3)
    doc.add_paragraph(f"Supply Detected: {'Yes' if rfp_info.get('scope_supply') else 'No'}")
    doc.add_paragraph(f"Installation Detected: {'Yes' if rfp_info.get('scope_installation') else 'No'}")
    doc.add_paragraph(f"Substitution Allowed: {'Yes' if rfp_info.get('scope_substitution_allowed') else 'No'}")
    doc.add_paragraph(f"No Substitution: {'Yes' if rfp_info.get('scope_no_substitution') else 'No'}")
    
    # Dates and Budgets
    if rfp_info.get('dates_found'):
        doc.add_paragraph(f"Dates Found: {', '.join(rfp_info['dates_found'])}")
    if rfp_info.get('budgets_found'):
        doc.add_paragraph(f"Budgets Found: {', '.join(rfp_info['budgets_found'])}")
    
    doc.add_page_break()
    
    # Preweights
    doc.add_heading("Criteria 1: Scope and Type Pre-Weighting (10%)", level=2)
    doc.add_paragraph(f"Detected Type: {preweight_result['detected_type']}")
    doc.add_paragraph(f"Condition: {preweight_result['condition']}")
    doc.add_paragraph(f"Detection Source: {preweight_result['detection_source']}")
    
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Company", "Pre-Weight Score"
    
    for company, weight in results.items():
        row = table.add_row().cells
        row[0].text = company
        row[1].text = f"{weight}/10"

    doc.save(file_path)

# ----------------- Streamlit App -------------------
def main():
    st.title("⚡ RFP Analyzer - Criteria 1 Pre-Weighting (NO API)")
    st.caption("Pure Text Extraction | Keyword-Based Analysis | No LLM/API Costs")
    
    # Show info banner
    st.success("""
    **✅ Zero API Costs:**
    - ✅ Text Extraction: PyMuPDF + docTR OCR (local, no API)
    - ✅ Analysis: Pure keyword matching (no AI calls)
    - ✅ Pre-Weighting: Rule-based logic (no LLM)
    - ✅ Export: Direct DOCX generation (no API)
    """)

    device = "cuda" if torch.cuda.is_available() else "cpu"
    doctr_model = load_doctr_model(device)

    # Show system status
    with st.expander("🔧 System Status"):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**GPU/Device Status:**")
            if torch.cuda.is_available():
                gpu_name = torch.cuda.get_device_name(0)
                gpu_memory = f"{torch.cuda.get_device_properties(0).total_memory / 1e9:.2f} GB"
                st.success(f"✅ GPU: {gpu_name}")
                st.info(f"💾 VRAM: {gpu_memory}")
            else:
                st.warning("⚠️ GPU: Not Available (Using CPU)")
            st.info(f"🖥️ Device: {device.upper()}")
        
        with col2:
            st.markdown("**Processing Mode:**")
            st.success("✅ Local Processing Only")
            st.success("✅ No API Calls")
            st.success("✅ No LLM Usage")
            st.info("🔧 OCR: docTR (GPU accelerated)")

    rfp_pdf = st.file_uploader("📄 Upload RFP PDF", type=["pdf"])

    if rfp_pdf and st.button("🚀 Analyze RFP (No API)"):
        rfp_filename = rfp_pdf.name
        
        # 1) Extract all text from PDF
        with st.spinner("📄 Extracting text from PDF (PyMuPDF + OCR fallback)..."):
            pages = extract_page_texts(rfp_pdf, doctr_model)
            full_text = "\n\n".join(pages)
        
        st.success(f"✅ Extracted {len(pages)} pages ({len(full_text):,} characters)")
        
        # 2) Extract basic info using keyword matching (NO API)
        with st.spinner("🔍 Analyzing content with keyword matching (no API)..."):
            rfp_info = extract_basic_info(full_text, rfp_filename)
            rfp_info['total_pages'] = len(pages)
        
        # Display extracted information
        st.markdown("---")
        st.subheader("📋 RFP Basic Information (Keyword-Based)")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Pages", rfp_info['total_pages'])
            st.metric("Total Characters", f"{rfp_info['total_chars']:,}")
        with col2:
            st.metric("Detected Type", rfp_info['project_type'])
            st.metric("Dates Found", len(rfp_info['dates_found']))
        with col3:
            st.metric("Budgets Found", len(rfp_info['budgets_found']))
        
        # Scope detection
        st.markdown("### 🎯 Scope Detection")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if rfp_info['scope_supply']:
                st.success("✅ Supply Detected")
            else:
                st.error("❌ Supply Not Found")
        with col2:
            if rfp_info['scope_installation']:
                st.success("✅ Installation Detected")
            else:
                st.error("❌ Installation Not Found")
        with col3:
            if rfp_info['scope_substitution_allowed']:
                st.success("✅ Substitution Allowed")
            else:
                st.warning("⚠️ Not Mentioned")
        with col4:
            if rfp_info['scope_no_substitution']:
                st.error("❌ No Substitution")
            else:
                st.info("ℹ️ Not Specified")
        
        # Show extracted dates and budgets
        with st.expander("📅 Extracted Dates"):
            if rfp_info['dates_found']:
                for date in rfp_info['dates_found']:
                    st.markdown(f"- {date}")
            else:
                st.write("No dates found")
        
        with st.expander("💰 Extracted Budgets"):
            if rfp_info['budgets_found']:
                for budget in rfp_info['budgets_found']:
                    st.markdown(f"- {budget}")
            else:
                st.write("No budget amounts found")
        
        # Show sample text
        with st.expander("📄 Sample Text (First 2000 characters)"):
            st.text(full_text[:2000] + "..." if len(full_text) > 2000 else full_text)
        
        # 3) Apply Criteria 1 Pre-Weighting (NO API)
        st.markdown("---")
        st.subheader("🎯 Criteria 1: Scope and Type Pre-Weighting (10%)")
        
        with st.spinner("⚖️ Applying pre-weighting rules (no API)..."):
            preweight_result = apply_scope_preweights(rfp_info, full_text)
        
        # Display detection results
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("**📊 Detected Project Type:**")
            detected_type = preweight_result["detected_type"]
            if "lighting" in detected_type.lower() or "led" in detected_type.lower():
                st.success(f"💡 {detected_type}")
            elif detected_type != "Unknown":
                st.info(f"🟢 {detected_type}")
            else:
                st.warning(f"⚪ {detected_type}")
        
        with col2:
            st.markdown("**📋 Specific Condition:**")
            condition = preweight_result["condition"]
            st.write(f"_{condition}_")
        
        with col3:
            st.markdown("**📍 Detection Source:**")
            detection_source = preweight_result.get("detection_source", "N/A")
            if detection_source == "Filename + Content":
                st.success(f"📄 {detection_source}")
            elif detection_source == "Content":
                st.info(f"📝 {detection_source}")
            else:
                st.warning(f"❓ {detection_source}")
        
        # Display pre-weights in a table
        st.markdown("**⚖️ 10% Baseline Pre-Weights:**")
        
        preweights = preweight_result['preweights']
        
        # Visual cards
        cols = st.columns(3)
        for idx, (company, weight) in enumerate(preweights.items()):
            with cols[idx]:
                if weight == 10:
                    st.success(f"**{company}**\n\n{weight}/10 ✅")
                elif weight > 0:
                    st.info(f"**{company}**\n\n{weight}/10 ⚠️")
                else:
                    st.error(f"**{company}**\n\n{weight}/10 ❌")
        
        # Comparison table
        st.markdown("### 📊 Pre-Weighting Comparison")
        preweights_df = pd.DataFrame([
            {
                "Company": company,
                "Pre-Weight Score": f"{weight}/10",
                "Percentage": f"{weight*10}%",
                "Status": "✅ Advantage" if weight > 0 else "⚪ Neutral"
            }
            for company, weight in preweights.items()
        ])
        
        st.dataframe(preweights_df, use_container_width=True, hide_index=True)
        
        # Explanation
        st.markdown("### 📝 Evaluation Logic Applied")
        if preweight_result["detected_type"] == "Lighting":
            st.markdown("""
            **Lighting Project Rules:**
            - **Supply + Substitution Allowed**: IKIO = 10, METCO = 0, SUNSPRINT = 0
            - **Supply + Installation + Substitution Allowed**: All = 10
            - **Supply + Installation + No Substitution**: IKIO = 0, METCO = 10, SUNSPRINT = 10
            - **Installation Only (No Supply)**: IKIO = 0, METCO = 10, SUNSPRINT = 10
            - **Default/Other Lighting Cases**: All = 10
            """)
        else:
            st.markdown("""
            **Non-Lighting Project Rules:**
            - **All Non-Lighting Types** (HVAC, Solar, Water, Wastewater, Building Envelope, ESCO, Energy Saving, Generator):
              - IKIO = 0
              - METCO = 10
              - SUNSPRINT = 10
            """)
        
        # Show detailed pre-weight logic
        with st.expander("📖 View Complete Pre-Weighting Rules"):
            st.markdown("#### 🟡 Lighting Project Rules")
            st.code("""
1. Supply + Substitution Allowed → IKIO=10%, others=0%
2. Supply + Installation + Substitution Allowed → All=10%
3. Supply + Installation + No Substitution → IKIO=0%, METCO=10%, SUNSPRINT=10%
4. Installation Only → IKIO=0%, METCO=10%, SUNSPRINT=10%
5. Default Lighting → All=10%
            """, language="text")
            
            st.markdown("#### 🟢 Non-Lighting Project Rules")
            st.code("""
• HVAC → IKIO=0%, METCO=10%, SUNSPRINT=10%
• Solar → IKIO=0%, METCO=10%, SUNSPRINT=10%
• Water/Waste Water → IKIO=0%, METCO=10%, SUNSPRINT=10%
• Building Envelope → IKIO=0%, METCO=10%, SUNSPRINT=10%
• ESCO → IKIO=0%, METCO=10%, SUNSPRINT=10%
• Emergency Generator → IKIO=0%, METCO=10%, SUNSPRINT=10%
            """, language="text")
            
            st.markdown("---")
            st.markdown(f"**Current Detection:**")
            st.write(f"• **Project Type:** {detected_type}")
            st.write(f"• **Condition:** {condition}")
            st.write(f"• **Detection Source:** {detection_source}")
            st.write(f"• **RFP Filename:** `{rfp_filename}`")
            st.write(f"• **Pre-Weights Applied:** {preweights}")
        
        # Export to DOCX
        st.markdown("---")
        output_file = "RFP_PreWeighting_Report.docx"
        export_to_docx(preweights, rfp_info, preweight_result, output_file)
        with open(output_file, "rb") as f:
            st.download_button(
                "⬇️ Download Pre-Weighting Report (DOCX)",
                f,
                file_name=output_file,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        st.success("✅ Analysis complete! No API calls were made - everything processed locally.")

if __name__ == "__main__":
    st.set_page_config(
        page_title="RFP Analyzer - No API",
        layout="wide",
        initial_sidebar_state="collapsed"
    )
    main()
