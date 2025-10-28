import streamlit as st
import os
import json
import torch
import pandas as pd
from doctr.models import ocr_predictor
from doctr.io import DocumentFile
from dotenv import load_dotenv
import requests
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer
import fitz
import re
from typing import List

# ----------------- Checklist Loader (Excel) -------------------
def load_checklist(file_path: str = "checklist.xlsx") -> pd.DataFrame:
    """Load and clean the checklist from Excel.

    Expects a worksheet named 'checklist' with columns:
    - Requirements
    - Criticality
    """
    df = pd.read_excel(file_path, sheet_name="checklist")
    df = df.rename(columns=lambda x: str(x).strip())
    required_cols = ["Requirements", "Criticality"]
    for col in required_cols:
        if col not in df.columns:
            raise ValueError(f"Missing column '{col}' in checklist.xlsx")
    df = df[["Requirements", "Criticality"]].dropna(subset=["Requirements"])  # keep only non-empty requirements
    return df

# ----------------- OCR -------------------
@st.cache_resource
def load_doctr_model(device: str = "cuda"):
    try:
        return ocr_predictor(pretrained=True).to(device)
    except Exception:
        return ocr_predictor(pretrained=True).to("cpu")

def extract_pages_with_doctr(pdf_bytes: bytes, model) -> List[str]:
    """Return a list of page-level texts using docTR OCR."""
    doc = DocumentFile.from_pdf(pdf_bytes)
    result = model(doc)
    return [page.render() for page in result.pages]

def extract_text_with_doctr(pdf_bytes: bytes, model) -> str:
    """Compatibility helper returning the whole document text joined by page breaks."""
    pages = extract_pages_with_doctr(pdf_bytes, model)
    return "\n\n--- Page Break ---\n\n".join(pages)

def extract_page_texts(file, doctr_model) -> List[str]:
    """
    Try to extract selectable text with PyMuPDF per-page.
    If no text is found (scanned PDF), OCR page-by-page with docTR.
    """
    pages = []
    try:
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                t = page.get_text("text")
                pages.append(t or "")
    except Exception:
        pages = []
    # If most pages are empty => OCR fallback
    if not pages or sum(1 for p in pages if p.strip()) < max(1, int(0.2 * len(pages))):
        file.seek(0)
        pdf_bytes = file.read()
        pages = extract_pages_with_doctr(pdf_bytes, doctr_model)
    file.seek(0)
    return pages

# ----------------- Ollama Client Helper -------------------
class OllamaClient:
    def __init__(self, base_url: str = "http://localhost:11434", model: str = "llama3"):
        self.base_url = base_url
        self.model = model
    
    def generate(self, prompt: str, json_mode: bool = True) -> str:
        url = f"{self.base_url}/api/generate"
        payload = {
            "model": self.model,
            "prompt": prompt,
            "stream": False,
            "format": "json" if json_mode else None,
        }
        try:
            resp = requests.post(url, json=payload, timeout=120)
            resp.raise_for_status()
            return resp.json().get("response", "{}")
        except Exception:
            return "{}"

# ----------------- File Reader (any type) -------------------
def read_company_file(file):
    try:
        if file.name.endswith(".txt"):
            return file.read().decode("utf-8", errors="ignore")
        elif file.name.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(file)
            return "\n".join([p.text for p in doc.paragraphs])
        elif file.name.endswith(".xlsx"):
            import pandas as pd
            df = pd.read_excel(file)
            return df.astype(str).to_string()
        elif file.name.endswith(".pdf"):
            pdf_bytes = file.read()
            model = load_doctr_model("cuda" if torch.cuda.is_available() else "cpu")
            # reuse page extractor for consistency
            pages = extract_pages_with_doctr(pdf_bytes, model)
            return "\n\n".join(pages)
        else:
            return file.read().decode("utf-8", errors="ignore")
    except Exception:
        return "Unsupported file format – please provide text."

# ----------------- Optimized Scorer -------------------
class OptimizedBidScorer:
    def __init__(self, client: OllamaClient):
        self.client = client

    def evaluate_company(self, rfp_summary: dict, company_name: str, company_context: str) -> dict:
        checklist = [
            {"requirement": "Full solicitation received?", "criticality": "C"},
            {"requirement": "Procurement requirements reviewed?", "criticality": "C"},
        ]
        checklist_text = "\n".join([f"- {c['requirement']} ({c['criticality']})" for c in checklist])

        prompt = f"""
You are a professional bid compliance analyst.
Evaluate the following company against the checklist.

RFP SUMMARY:
{rfp_summary}

COMPANY CONTEXT:
{company_context}

CHECKLIST:
{checklist_text}

Rules:
- For each requirement, return Yes / No / Partial with short justification.
- Output ONLY valid JSON with this exact structure:
{{
 "complianceAssessment":[
   {{"requirement":"...", "criticality":"C/NC", "evaluation":"Yes/No/Partial", "comments":"..."}}
 ]
}}
"""
        try:
            response_text = self.client.generate(prompt, json_mode=True)
            result = json.loads(response_text)
            assessments = result.get("complianceAssessment", [])
        except Exception:
            assessments = [{"requirement": c["requirement"], "criticality": c["criticality"],
                            "evaluation": "Partial", "comments": "Evaluation failed"} for c in checklist]

        total_score = 0
        for a in assessments:
            a["score"] = 1 if a.get("evaluation") == "Yes" else 0.5 if a.get("evaluation") == "Partial" else 0
            total_score += a["score"]

        return {
            "company": company_name,
            "complianceAssessment": assessments,
            "finalScoring": {
                "totalScore": total_score,
                "overallCompliancePercentage": round((total_score / max(1, len(checklist))) * 100, 1),
                "recommendation": "TBD",
                "keyRiskFactors": [],
                "actionItems": [],
            },
        }
# ----------------- Simple Text Analysis (NO API) -------------------
def extract_basic_info(text: str, filename: str) -> dict:
    """
    Extract basic information from RFP text using keyword matching.
    NO API calls - pure text analysis.
    """
    text_lower = text.lower()
    
    # Extract dates (simple regex)
    date_patterns = [
        r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',  # MM/DD/YYYY or DD-MM-YYYY
        r'\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2},?\s+\d{4}\b'
    ]
    dates = []
    for pattern in date_patterns:
        dates.extend(re.findall(pattern, text, re.IGNORECASE))
    
    # Extract potential costs/budgets
    budget_pattern = r'\$[\d,]+(?:\.\d{2})?'
    budgets = re.findall(budget_pattern, text)
    
    # Detect project type keywords
    project_types = []
    type_keywords = {
        'lighting': ['lighting', 'led', 'illumination', 'fixture', 'luminaire', 'lamp'],
        'hvac': ['hvac', 'heating', 'ventilation', 'air conditioning', 'cooling'],
        'solar': ['solar', 'photovoltaic', 'pv system'],
        'water': ['water', 'plumbing', 'water management'],
        'wastewater': ['wastewater', 'waste water', 'sewage'],
        'building envelope': ['building envelope', 'roofing', 'insulation', 'windows'],
        'esco': ['esco', 'energy service', 'performance contract'],
        'generator': ['generator', 'emergency power', 'backup power']
    }
    
    for proj_type, keywords in type_keywords.items():
        if any(kw in text_lower or kw in filename.lower() for kw in keywords):
            project_types.append(proj_type.upper())
    
    # Detect scope keywords
    has_supply = any(kw in text_lower for kw in ['supply', 'furnish', 'provide', 'material'])
    has_installation = any(kw in text_lower for kw in ['install', 'installation', 'labor'])
    has_substitution = any(kw in text_lower for kw in ['substitution allowed', 'or equal', 'approved equal', 'or equivalent'])
    has_no_substitution = any(kw in text_lower for kw in ['no substitution', 'as specified', 'exact match'])
    
    return {
        'project_type': ', '.join(project_types) if project_types else 'General/Unknown',
        'dates_found': dates[:10],  # First 10 dates
        'budgets_found': budgets[:10],  # First 10 budget mentions
        'scope_supply': has_supply,
        'scope_installation': has_installation,
        'scope_substitution_allowed': has_substitution,
        'scope_no_substitution': has_no_substitution,
        'total_pages': 0,  # Will be set later
        'total_chars': len(text),
        'filename': filename
    }

# ----------------- Criteria 1: Scope and Type Pre-Weighting (10%) -------------------
def apply_scope_preweights(rfp_info: dict, full_text: str) -> dict:
    """
    Apply Criteria 1 pre-weighting (10%) dynamically based on detected project type
    and scope. NO API - pure keyword matching.
    """
    LIGHTING_KEYWORDS = ["lighting", "led", "illumination", "fixture", "luminaire"]
    PROJECT_TYPE_KEYWORDS = {
        "hvac": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "solar": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "water": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "waste water": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "wastewater": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "building envelope": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "esco": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "energy saving": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
        "generator": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
    }

    # Combine filename + text for detection
    filename_part = rfp_info.get('filename', '').lower()
    text = (filename_part + " " + full_text).lower()

    preweights = {"IKIO": 0, "METCO": 0, "SUNSPRINT": 0}
    detected_type, condition = "Unknown", "N/A"
    detection_source = "N/A"

    # --- Lighting branch ---
    if any(k in text for k in LIGHTING_KEYWORDS):
        detected_type = "Lighting"
        has_supply = rfp_info.get('scope_supply', False)
        has_install = rfp_info.get('scope_installation', False)
        substitution_allowed = rfp_info.get('scope_substitution_allowed', False)
        no_substitution = rfp_info.get('scope_no_substitution', False)

        # Determine detection source
        if any(kw in filename_part for kw in LIGHTING_KEYWORDS):
            detection_source = "Filename + Content"
        else:
            detection_source = "Content"

        if has_supply and not has_install and substitution_allowed:
            preweights = {"IKIO": 10, "METCO": 0, "SUNSPRINT": 0}
            condition = "Supply + Substitution Allowed"
        elif has_supply and has_install and substitution_allowed:
            preweights = {"IKIO": 10, "METCO": 10, "SUNSPRINT": 10}
            condition = "Supply + Installation + Substitution Allowed"
        elif has_supply and has_install and no_substitution:
            preweights = {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10}
            condition = "Supply + Installation + Substitution Not Allowed"
        elif not has_supply and has_install:
            preweights = {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10}
            condition = "Installation Only"
        else:
            preweights = {"IKIO": 10, "METCO": 0, "SUNSPRINT": 0}
            condition = "Lighting (Default All)"
    else:
        # --- Non-lighting dynamic detection ---
        for kw, weights in PROJECT_TYPE_KEYWORDS.items():
            if kw in text:
                detected_type = kw.title()
                preweights = weights
                condition = f"Detected via keyword '{kw}'"
                
                # Determine detection source
                if kw in filename_part:
                    detection_source = "Filename + Content"
                else:
                    detection_source = "Content"
                break

    return {
        "detected_type": detected_type,
        "condition": condition,
        "preweights": preweights,
        "detection_source": detection_source
    }

# ----------------- Criteria 1 via API (minimal tokens) -------------------
def apply_scope_preweights_via_api(rfp_summary: dict, rfp_filename: str | None = None, raw_text: str | None = None) -> dict:
    """
    Use an external API (OpenAI-compatible chat.completions) to classify Criteria 1
    with a tiny JSON. Falls back to local keyword rules when API is not configured.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    base_url = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1/chat/completions")
    model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

    # Build compact text from summary to keep token usage low
    compact = (
        f"project_type: {rfp_summary.get('project_type','')}; "
        f"scope: {rfp_summary.get('scope','')}; "
        f"summary: {rfp_summary.get('summary','')}; "
    )[:6000]

    # Local heuristic flags from full raw text (helps when API misses cues)
    lt = (raw_text or compact).lower()
    local_has_supply = any(k in lt for k in ["supply", "furnish", "provide", "procure", "deliver", "material"])
    local_has_install = any(k in lt for k in [
        "install", "installation", "installing", "installed", "installment",  # include common misspelling
        "retrofit", "replacement", "replace", "erection"
    ])
    local_sub_allowed = any(k in lt for k in ["substitution allowed", "approved equal", "or equal", "or equivalent", "approved equivalent"])
    local_no_sub = any(k in lt for k in ["no substitution", "no substitutions", "no alternates", "no equals", "as specified only"]) 

    if not api_key:
        # Fallback to local rules using compact text as both scope and summary
        return apply_scope_preweights({
            'filename': rfp_filename or '',
            'scope_supply': ('supply' in compact.lower() or 'furnish' in compact.lower() or 'provide' in compact.lower()),
            'scope_installation': ('install' in compact.lower() or 'installation' in compact.lower()),
            'scope_substitution_allowed': ('substitution allowed' in compact.lower() or 'approved equal' in compact.lower() or 'or equal' in compact.lower()),
            'scope_no_substitution': ('no substitution' in compact.lower() or 'as specified only' in compact.lower() or 'exact match' in compact.lower()),
            'project_type': rfp_summary.get('project_type','')
        }, compact)

    prompt = f"""
Return ONLY JSON with fields:
{{
  "detected_type": "Lighting|HVAC|Solar|Water|Waste Water|Building Envelope|ESCO|Energy Saving|Generator|Unknown",
  "has_supply": true|false,
  "has_install": true|false,
  "substitution_allowed": true|false,
  "no_substitution": true|false
}}
Use the filename (if provided) and this compact RFP summary to classify type and scope.
FILENAME: {rfp_filename or 'N/A'}
TEXT: {compact}
"""
    try:
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        body = {
            "model": model,
            "temperature": 0,
            "response_format": {"type": "json_object"},
            "messages": [{"role": "user", "content": prompt}],
        }
        resp = requests.post(base_url, headers=headers, json=body, timeout=60)
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
        data = json.loads(content)
        dtype = str(data.get("detected_type", "Unknown")).strip()
        has_supply = bool(data.get("has_supply", False)) or local_has_supply
        has_install = bool(data.get("has_install", False)) or local_has_install
        sub_allowed = bool(data.get("substitution_allowed", False)) or local_sub_allowed
        no_sub = bool(data.get("no_substitution", False)) or local_no_sub

        # Map to preweights
        if dtype.lower() == "lighting":
            if has_supply and not has_install and sub_allowed:
                preweights = {"IKIO": 10, "METCO": 0, "SUNSPRINT": 0}
                condition = "Supply + Substitution Allowed"
            elif has_supply and has_install and sub_allowed:
                preweights = {"IKIO": 10, "METCO": 10, "SUNSPRINT": 10}
                condition = "Supply + Installation + Substitution Allowed"
            elif has_supply and has_install and no_sub:
                preweights = {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10}
                condition = "Supply + Installation + Substitution Not Allowed"
            elif not has_supply and has_install:
                preweights = {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10}
                condition = "Installation Only"
            else:
                preweights = {"IKIO": 10, "METCO": 0, "SUNSPRINT": 0}
                condition = "Lighting (Default All)"
            return {"detected_type": "Lighting", "condition": condition, "preweights": preweights, "detection_source": "API"}
        else:
            mapping = {
                "hvac": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
                "solar": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
                "water": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
                "waste water": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
                "building envelope": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
                "esco": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
                "energy saving": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
                "generator": {"IKIO": 0, "METCO": 10, "SUNSPRINT": 10},
            }
            key = dtype.lower()
            if key in mapping:
                return {"detected_type": dtype.title(), "condition": f"Detected via API as {dtype}", "preweights": mapping[key], "detection_source": "API"}
    except Exception:
        pass

    # Fallback to local rules
    return apply_scope_preweights({
        'filename': rfp_filename or '',
        'scope_supply': ('supply' in compact.lower() or 'furnish' in compact.lower() or 'provide' in compact.lower()),
        'scope_installation': ('install' in compact.lower() or 'installation' in compact.lower()),
        'scope_substitution_allowed': ('substitution allowed' in compact.lower() or 'approved equal' in compact.lower() or 'or equal' in compact.lower()),
        'scope_no_substitution': ('no substitution' in compact.lower() or 'as specified only' in compact.lower() or 'exact match' in compact.lower()),
        'project_type': rfp_summary.get('project_type','')
    }, compact)

# ----------------- Criteria 2: Project Location (10%) -------------------
def apply_location_preweights(raw_text: str, rfp_summary: dict, rfp_filename: str | None = None) -> dict:
    """
    Detect project location and give 10% ONLY if the project site matches a registered location:
      - IKIO: Indiana / Indianapolis
      - SUNSPRINT: Indiana / Indianapolis
      - METCO: Dallas / Texas

    Strategy to reduce false positives:
      - Search within windows around location anchors ("project location", "located in", "site location", etc.)
      - Exclude contexts like vendor registration/licensing lines
      - Use whole-word boundary regexes
    """
    # Build text pool (filename + ollama fields + raw)
    text_full = "\n".join([
        rfp_filename or "",
        rfp_summary.get("scope", "") or "",
        rfp_summary.get("summary", "") or "",
        raw_text or ""
    ])
    t = text_full.lower()

    # Early presence of other states to avoid accidental matches
    other_states = ["missouri", " mo ", "kansas", " il ", "illinois", "ohio", "kentucky"]

    # Anchors around which we trust location mentions
    anchors = [
        "project location", "project is located", "site location", "location:", "city:", "state:",
        "located at", "located in", "place of performance", "place of work", "project address", "site address"
    ]

    # Helper to collect windows around anchors
    windows = []
    for a in anchors:
        start = 0
        while True:
            idx = t.find(a, start)
            if idx == -1:
                break
            win = t[max(0, idx-100): idx+200]
            windows.append(win)
            start = idx + len(a)

    # If no anchors found, fall back to a trimmed global window to avoid full-document noise
    if not windows:
        windows = [t[:3000]]

    def any_regex(patterns: list, hay: str) -> bool:
        import re as _re
        return any(_re.search(p, hay) is not None for p in patterns)

    # Strict word-boundary patterns
    in_patterns = [r"\bindiana\b", r"\bindianapolis\b", r"\bindianpolis\b", r"indian\-polis"]
    tx_patterns = [r"\bdallas\b", r"\btexas\b", r"\btx\b"]

    # Exclusions for registration/licensing text
    exclude_patterns = ["registered", "licens", "vendor", "bidder", "mailing", "office", "headquarter"]

    def window_valid(win: str) -> bool:
        return not any(ex in win for ex in exclude_patterns)

    # Track evidence snippets
    evidence_snippets = []
    import re as _re
    for w in windows:
        if not window_valid(w):
            continue
        if any_regex(in_patterns, w) or any_regex(tx_patterns, w):
            # Clean and compress whitespace for display
            snippet = _re.sub(r"\s+", " ", w.strip())
            evidence_snippets.append(snippet[:240])

    found_indiana = any(window_valid(w) and any_regex(in_patterns, w) for w in windows)
    found_dallas = any(window_valid(w) and any_regex(tx_patterns, w) for w in windows)

    # If another state is explicitly present and no anchor-confirmed Indiana/Dallas, avoid awarding
    if any(s in t for s in other_states):
        if not found_indiana:
            found_indiana = False
        if not found_dallas:
            found_dallas = False

    pre = {"IKIO": 0, "METCO": 0, "SUNSPRINT": 0}
    matched = []
    location_names = []
    if found_indiana:
        pre["IKIO"] = 10
        pre["SUNSPRINT"] = 10
        matched.append("Indiana/Indianapolis")
        # Prefer city when present
        if any_regex([r"\bindianapolis\b", r"\bindianpolis\b"], t):
            location_names.append("Indianapolis, IN")
        else:
            location_names.append("Indiana")
    if found_dallas:
        pre["METCO"] = 10
        matched.append("Dallas/Texas")
        if any_regex([r"\bdallas\b"], t):
            location_names.append("Dallas, TX")
        else:
            location_names.append("Texas")

    # If no registered location matched, still try to display a friendly project location name
    if not location_names:
        # Simple state detection to show user-friendly location while keeping scores at 0
        state_map = {
            "indiana": "Indiana",
            "texas": "Texas",
            "missouri": "Missouri",
            "illinois": "Illinois",
            "ohio": "Ohio",
            "kentucky": "Kentucky",
            "kansas": "Kansas"
        }
        chosen_state = None
        for w in windows:
            if not window_valid(w):
                continue
            lw = w.lower()
            for key, name in state_map.items():
                if re.search(r"\\b" + key + r"\\b", lw):
                    chosen_state = name
                    snippet = re.sub(r"\\s+", " ", w.strip())
                    evidence_snippets.append(snippet[:240])
                    break
            if chosen_state:
                break
        if chosen_state:
            location_names.append(chosen_state)

    return {
        "preweights": pre,
        # All locations we detected from the RFP text (for display)
        "detected_locations": location_names or ["Unknown"],
        # Only registered locations that matched (these drive scoring)
        "scored_locations": matched,
        "location_display": " / ".join(location_names) if location_names else "Unknown",
        "detection_source": "Anchored Location Detection",
        "evidence_snippets": evidence_snippets[:5]
    }
# ----------------- DOCX Export -------------------
def export_to_docx(results: dict, rfp_info_or_summary: dict, file_path: str, preweight_result: dict = None):
    """Flexible DOCX export supporting both old (summary-only) and new (preweight) flows."""
    doc = DocxDocument()
    doc.add_heading("Bid Compliance Evaluation Report", level=1)

    # If preweight info provided, include basic RFP info section
    if preweight_result is not None:
        doc.add_heading("RFP Information", level=2)
        doc.add_paragraph(f"Filename: {rfp_info_or_summary.get('filename','N/A')}")
        doc.add_paragraph(f"Total Pages: {rfp_info_or_summary.get('total_pages',0)}")
        doc.add_paragraph(f"Total Characters: {rfp_info_or_summary.get('total_chars',0)}")
        doc.add_paragraph(f"Detected Project Type: {rfp_info_or_summary.get('project_type','N/A')}")
        doc.add_heading("Criteria 1 Pre-Weighting (10%)", level=2)
        doc.add_paragraph(f"Detected Type: {preweight_result.get('detected_type','N/A')}")
        doc.add_paragraph(f"Condition: {preweight_result.get('condition','N/A')}")
        doc.add_paragraph(f"Detection Source: {preweight_result.get('detection_source','N/A')}")
        tbl = doc.add_table(rows=1, cols=2)
        hdr = tbl.rows[0].cells
        hdr[0].text, hdr[1].text = "Company", "Pre-Weight Score"
        for company, weight in results.items():
            row = tbl.add_row().cells
            row[0].text = company
            row[1].text = f"{weight}/10"

    # Minimal compatibility section when used with old flow
    if preweight_result is None:
        doc.add_heading("RFP Summary", level=2)
        doc.add_paragraph(f"Project Type: {rfp_info_or_summary.get('project_type','N/A')}")
        doc.add_paragraph(f"Due Date: {rfp_info_or_summary.get('due_date','N/A')}")

    doc.save(file_path)

# ----------------- Streamlit App -------------------
def main():
    st.title("⚡ Stable Bid Analyzer (Ollama + DOCX Export)")

    device = "cuda" if torch.cuda.is_available() else "cpu"
    doctr_model = load_doctr_model(device)

    load_dotenv()
    # Initialize Ollama client (default: llama3 model on localhost:11434)
    ollama_url = os.getenv("OLLAMA_URL", "http://localhost:11434")
    ollama_model = os.getenv("OLLAMA_MODEL", "llama3")
    client = OllamaClient(base_url=ollama_url, model=ollama_model)
    scorer = OptimizedBidScorer(client)

    rfp_pdf = st.file_uploader("📄 Upload RFP PDF", type=["pdf"])
    st.subheader("📂 Upload Company Contexts (any file)")
    col1, col2, col3 = st.columns(3)
    with col1: ikio_file = st.file_uploader("IKIO", type=None, key="ikio")
    with col2: metco_file = st.file_uploader("METCO", type=None, key="metco")
    with col3: sunsprint_file = st.file_uploader("Sunsprint", type=None, key="sunsprint")


    if rfp_pdf and st.button("🚀 Evaluate All Companies"):
        rfp_bytes = rfp_pdf.read()
        with st.spinner("Extracting RFP text..."):
            rfp_text = extract_text_with_doctr(rfp_bytes, doctr_model)
        
        # Generate RFP summary using Ollama
        with st.spinner("Analyzing RFP with Ollama..."):
            sanitized = rfp_text.encode('ascii', 'ignore').decode('ascii')
            rfp_prompt = f"""Analyze this RFP and extract key information. Return ONLY valid JSON:
{{
 "project_type": "brief project type",
 "scope": "brief project scope",
 "key_requirements": ["requirement1", "requirement2"],
 "due_date": "extracted due date if found, otherwise N/A",
 "total_cost": "estimated cost if mentioned, otherwise N/A"
}}

RFP Text (first 8000 chars):
{sanitized[:8000]}"""
            
            try:
                rfp_response = client.generate(rfp_prompt, json_mode=True)
                rfp_summary = json.loads(rfp_response)
                st.success("✅ RFP analyzed successfully")
            except Exception as e:
                st.warning(f"RFP analysis error: {str(e)}, using basic summary")
                rfp_summary = {"project_type": "General", "scope": rfp_text[:1500], "due_date": "N/A"}

        # ---- Criteria 1: API-first classification (fallback to local) ----
        with st.spinner("Applying Criteria 1 (Type & Scope) via API..."):
            pre1 = apply_scope_preweights_via_api(rfp_summary, rfp_pdf.name, raw_text=rfp_text)

        st.subheader("🎯 Criteria 1: Scope & Type Pre-Weighting (10%)")
        cols_pre = st.columns(3)
        with cols_pre[0]:
            st.metric("Detected Type", pre1.get("detected_type", "Unknown"))
        with cols_pre[1]:
            st.metric("Condition", pre1.get("condition", "N/A"))
        with cols_pre[2]:
            st.metric("Source", pre1.get("detection_source", "N/A"))

        pw = pre1.get("preweights", {"IKIO": 0, "METCO": 0, "SUNSPRINT": 0})
        st.table(pd.DataFrame([
            {"Company": c, "Pre-Weight": f"{w}/10", "Percent": f"{w*10}%"}
            for c, w in pw.items()
        ]))

        # Build basic info for DOCX export
        page_sep = "--- Page Break ---"
        total_pages = len(rfp_text.split(page_sep)) if page_sep in rfp_text else 0
        rfp_info = {
            "filename": rfp_pdf.name,
            "total_pages": total_pages,
            "total_chars": len(rfp_text),
            "project_type": rfp_summary.get("project_type", "N/A"),
        }

        # ---- Criteria 2: Project Location (10%) ----
        st.subheader("📍 Criteria 2: Project Location (10%)")
        pre2 = apply_location_preweights(rfp_text, rfp_summary, rfp_pdf.name)
        loc_pw = pre2.get("preweights", {"IKIO": 0, "METCO": 0, "SUNSPRINT": 0})
        st.caption(f"Detected: {pre2.get('location_display', 'Unknown')}")
        ev = pre2.get('evidence_snippets', [])
        if ev:
            with st.expander("Show location evidence"):
                for i, s in enumerate(ev, start=1):
                    st.markdown(f"{i}. {s}")
        st.table(pd.DataFrame([
            {"Company": c, "Pre-Weight": f"{w}/10", "Percent": f"{w*10}%"}
            for c, w in loc_pw.items()
        ]))

        # ---- Combined Pre-Weight (Criteria 1 + Criteria 2) ----
        st.subheader("🧮 Combined Pre-Weight (C1 + C2)")
        c1 = pw
        c2 = loc_pw
        companies_list = ["IKIO", "METCO", "SUNSPRINT"]
        combined = {c: int(c1.get(c, 0)) + int(c2.get(c, 0)) for c in companies_list}
        st.table(pd.DataFrame([
            {"Company": c, "Criteria 1": f"{c1.get(c,0)}/10", "Criteria 2": f"{c2.get(c,0)}/10", "Total": f"{combined[c]}/20"}
            for c in companies_list
        ]))
        selected_for_eval = sorted(companies_list, key=lambda x: combined[x], reverse=True)[:2]
        st.info(f"Top 2 by pre-weight: {', '.join(selected_for_eval)}")

        # ---- Show Ollama-extracted text (what was extracted) ----
        st.markdown("---")
        st.subheader("📝 Ollama Extracted Text")
        if rfp_summary.get("scope"):
            st.markdown("**Scope (from Ollama):**")
            st.info(rfp_summary.get("scope", ""))
        if rfp_summary.get("summary"):
            st.markdown("**Summary (from Ollama):**")
            st.write(rfp_summary.get("summary", ""))
        with st.expander("View full Ollama JSON"):
            st.json(rfp_summary)

        # ---- Evaluate only the best two companies (checklist.xlsx) ----
        company_files = {"IKIO": ikio_file, "METCO": metco_file, "SUNSPRINT": sunsprint_file}
        results = {}
        # Load Excel exactly and preserve its column order for display
        checklist_df = load_checklist("checklist.xlsx")
        checklist_items = [
            {"requirement": row.Requirements, "criticality": row.Criticality}
            for row in checklist_df.itertuples(index=False)
        ]

        for name in selected_for_eval:
            f = company_files.get(name)
            ctx = ""
            if f:
                ctx = read_company_file(f)
                save_company_context(name, ctx)
            else:
                cached = load_company_context(name)
                if cached:
                    ctx = cached
                else:
                    st.warning(f"No context for {name}; skipping evaluation.")
                    continue
            with st.spinner(f"Evaluating {name}..."):
                eval_items = evaluate_checklist_with_api(name, rfp_summary, ctx, checklist_items, rfp_full_text=rfp_text)
                total_score = sum(it["score"] for it in eval_items)
                max_score = len(eval_items) * 1.0
                compliance_pct = round((total_score / max_score) * 100, 1) if max_score else 0
                results[name] = {
                    "complianceAssessment": eval_items,
                    "finalScoring": {
                        "totalScore": round(total_score, 2),
                        "overallCompliancePercentage": compliance_pct,
                        "recommendation": "Go" if compliance_pct >= 85 else "Conditional Go" if compliance_pct >= 70 else "No-Go",
                    },
                }

        for name, res in results.items():
            st.subheader(f"🏢 {name} – Evaluation")
            fs = res["finalScoring"]
            st.metric("Total Score", fs["totalScore"])
            st.metric("Compliance %", fs["overallCompliancePercentage"])
            st.metric("Recommendation", fs["recommendation"])
            # Build a table that matches the Excel-like header shown in the image
            eval_df = pd.DataFrame(res["complianceAssessment"]).rename(columns={
                "requirement": "Requirements",
                "criticality": "Criticality",
                "evaluation": "Evaluation",
                "suggestions": "Suggestions",
                "rfp_context": "Context/Answers from RFP",
                "company_context": "Context/Answers from Company Documents",
                "score": "Score",
                "comments": "Comments",
            })

            # Prepare base columns in desired order
            display_cols = [
                "Sr. No.",
                "Requirements",
                "For Comparing Context Source",
                "Context/Answers from RFP",
                "Context/Answers from Company Documents",
                "Suggestions on comparing both RFP+Company Documents",
                "Criticality",
                "Evaluation",
                "Comments",
                "Score",
            ]

            # Derive Sr. No. and mapping columns
            if not eval_df.empty:
                eval_df = eval_df.reset_index(drop=True)
                eval_df["Sr. No."] = eval_df.index + 1
                eval_df["For Comparing Context Source"] = "RFP + Company"
                eval_df["Suggestions on comparing both RFP+Company Documents"] = eval_df["Suggestions"].fillna("")
                # Reorder and fill missing columns
                for col in display_cols:
                    if col not in eval_df.columns:
                        eval_df[col] = ""
                st.dataframe(eval_df[display_cols], use_container_width=True)
            else:
                base_df = checklist_df.copy().reset_index(drop=True)
                base_df["Sr. No."] = base_df.index + 1
                base_df["For Comparing Context Source"] = "RFP + Company"
                base_df["Context/Answers from RFP"] = ""
                base_df["Context/Answers from Company Documents"] = ""
                base_df["Suggestions on comparing both RFP+Company Documents"] = ""
                base_df["Evaluation"] = ""
                base_df["Comments"] = ""
                base_df["Score"] = 0
                st.dataframe(base_df[[
                    "Sr. No.", "Requirements", "For Comparing Context Source",
                    "Context/Answers from RFP", "Context/Answers from Company Documents",
                    "Suggestions on comparing both RFP+Company Documents",
                    "Criticality", "Evaluation", "Comments", "Score"
                ]], use_container_width=True)

        if results:
            comp_df = pd.DataFrame([
                {
                    "Company": n,
                    "Score": r["finalScoring"]["totalScore"],
                    "Compliance %": r["finalScoring"]["overallCompliancePercentage"],
                    "Recommendation": r["finalScoring"]["recommendation"],
                }
                for n, r in results.items()
            ])
            st.subheader("📊 Company Comparison")
            st.table(comp_df)

# ----------------- Simple Company Context Cache -------------------
def get_company_cache_path(company: str) -> str:
    cache_dir = "company_context_cache"
    os.makedirs(cache_dir, exist_ok=True)
    return os.path.join(cache_dir, f"{company}.txt")

def save_company_context(company: str, text: str) -> None:
    try:
        with open(get_company_cache_path(company), "w", encoding="utf-8") as f:
            f.write(text or "")
    except Exception:
        pass

def load_company_context(company: str) -> str:
    try:
        with open(get_company_cache_path(company), "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return ""

# ----------------- Checklist Loader -------------------
def load_checklist_items(excel_path: str = "checklist.xlsx") -> List[dict]:
    """Load checklist items from an Excel file.
    Expected columns (case-insensitive): requirement, criticality.
    If criticality missing, default to 'C'.
    """
    try:
        import pandas as _pd
        df = _pd.read_excel(excel_path)
        cols = {c.strip().lower(): c for c in df.columns}
        req_col = cols.get("requirement") or cols.get("requirements") or list(df.columns)[0]
        crit_col = cols.get("criticality") if "criticality" in cols else None
        items = []
        for _, row in df.iterrows():
            requirement = str(row.get(req_col, "")).strip()
            if not requirement:
                continue
            criticality = str(row.get(crit_col, "C")).strip().upper() if crit_col else "C"
            items.append({"requirement": requirement, "criticality": criticality or "C"})
        return items
    except Exception:
        # Fallback minimal items if file not present
        return [
            {"requirement": "Scope alignment with RFP", "criticality": "C"},
            {"requirement": "Schedule feasibility", "criticality": "C"},
        ]

# ----------------- Compact RFP context -------------------
def build_compact_rfp_context(rfp_summary: dict, max_chars: int = 3000) -> str:
    parts = []
    parts.append(f"Project Type: {rfp_summary.get('project_type','')}")
    if rfp_summary.get("scope"):
        parts.append(f"Scope: {rfp_summary.get('scope','')}")
    if rfp_summary.get("summary"):
        parts.append(f"Summary: {rfp_summary.get('summary','')}")
    for key in ["key_requirements", "technical_requirements", "bid_requirements"]:
        vals = rfp_summary.get(key) or []
        if isinstance(vals, list) and vals:
            parts.append(f"{key}: {', '.join(vals[:20])}")
    s = "\n".join(parts)
    return s[:max_chars]

# ----------------- Lightweight snippet retriever -------------------
def find_best_snippet(text: str, requirement: str, max_len: int = 180) -> str:
    """Return a short snippet from text that best matches the requirement.
    Uses simple token overlap over sentence-like chunks to avoid heavy dependencies.
    """
    if not text or not requirement:
        return ""
    try:
        import re as _re
        # Normalize and split text into chunks (sentences/lines)
        cleaned = _re.sub(r"\s+", " ", str(text))
        # Split by punctuation and newlines
        parts = _re.split(r"(?<=[\.!?])\s+|\n+|;\s+", cleaned)
        # Prepare tokens
        stop = {"the","a","an","and","or","of","to","for","in","on","at","by","with","is","are","be","as","that","this","these","those","from"}
        req_tokens = [t for t in _re.findall(r"[a-z0-9]+", requirement.lower()) if t not in stop]
        if not req_tokens:
            req_tokens = _re.findall(r"[a-z0-9]+", requirement.lower())
        best_score, best = 0, ""
        for p in parts:
            ptoks = [t for t in _re.findall(r"[a-z0-9]+", p.lower()) if t not in stop]
            if not ptoks:
                continue
            overlap = len(set(req_tokens) & set(ptoks))
            if overlap > best_score:
                best_score, best = overlap, p.strip()
        return (best or cleaned)[:max_len]
    except Exception:
        return str(text)[:max_len]

# ----------------- API Checklist Evaluation -------------------
def evaluate_checklist_with_api(company_name: str, rfp_summary: dict, company_text: str, items: List[dict], rfp_full_text: str = "") -> List[dict]:
    """Call OpenAI-compatible API once per company to evaluate all checklist items.
    Returns list of dicts with requirement, criticality, evaluation, comments, score.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    base_url = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1/chat/completions")
    model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

    # Build compact context
    rfp_ctx = build_compact_rfp_context(rfp_summary)
    # Trim company text to keep tokens small
    company_ctx = (company_text or "")[:4000]

    # Prepare items json for prompt with short evidence hints
    items_text = []
    for it in items:
        req = it["requirement"]
        crit = (it.get("criticality") or "C")
        rfp_hint = find_best_snippet(rfp_full_text, req, max_len=180)
        company_hint = find_best_snippet(company_text, req, max_len=180)
        items_text.append({
            "requirement": req,
            "criticality": crit,
            "rfp_hint": rfp_hint,
            "company_hint": company_hint,
        })

    prompt = (
        "You are evaluating a company against an RFP checklist."
        " Return ONLY JSON with an array 'results' where each item is {requirement, criticality, evaluation, suggestions, rfp_context, company_context, comments}.\n"
        "- evaluation must be one of Yes|Partial|No.\n"
        "- rfp_context: brief evidence from RFP (<=25 words).\n"
        "- company_context: brief evidence from company doc (<=25 words).\n"
        "- suggestions: brief action/observation (<=20 words).\n\n"
        f"RFP Context (compact):\n{rfp_ctx}\n\n"
        f"Company Context (trimmed):\n{company_ctx}\n\n"
        f"Checklist Items with HINTS (use hints as primary evidence):\n{json.dumps(items_text, ensure_ascii=False)}\n\n"
        "Output JSON schema: {\"results\":[{\"requirement\":\"...\",\"criticality\":\"C/NC\",\"evaluation\":\"Yes|Partial|No\",\"suggestions\":\"...\",\"rfp_context\":\"...\",\"company_context\":\"...\",\"comments\":\"...\"}]}"
    )

    results = []
    try:
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        body = {
            "model": model,
            "temperature": 0,
            "response_format": {"type": "json_object"},
            "messages": [{"role": "user", "content": prompt}],
        }
        resp = requests.post(base_url, headers=headers, json=body, timeout=90)
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
        data = json.loads(content)
        for it in data.get("results", []):
            requirement = it.get("requirement", "")
            criticality = it.get("criticality", "C")
            evaluation = it.get("evaluation", "Partial")
            comments = it.get("suggestions", "")
            rfp_ctx_item = it.get("rfp_context", "")
            company_ctx_item = it.get("company_context", "")
            extra_comments = it.get("comments", "")
            # scoring: Yes=2, Partial=1, No=0
            score = 2 if evaluation == "Yes" else 1 if evaluation == "Partial" else 0
            results.append({
                "requirement": requirement,
                "criticality": criticality,
                "evaluation": evaluation,
                "suggestions": comments,
                "rfp_context": rfp_ctx_item,
                "company_context": company_ctx_item,
                "comments": extra_comments,
                "score": score,
            })
    except Exception:
        # Fallback: mark all as Partial
        for it in items:
            results.append({
                "requirement": it["requirement"],
                "criticality": it.get("criticality", "C"),
                "evaluation": "Partial",
                "suggestions": "Evaluation failed",
                "rfp_context": "",
                "company_context": "",
                "comments": "",
                "score": 1,
            })
    return results

if __name__ == "__main__":
    st.set_page_config(page_title="Stable Bid Evaluator Optimized", layout="wide")
    main()
