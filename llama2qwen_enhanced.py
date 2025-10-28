# rfp_analyzer_batched_ollama_llama3_10pages_enhanced.py
import streamlit as st
import fitz
import os
import json
import numpy as np
import torch
import pandas as pd
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv
from docx import Document
from doctr.models import ocr_predictor
from doctr.io import DocumentFile
import re
from typing import List, Dict, Any

# === New: Ollama (Llama 3) ===
import ollama

# ========================== Config ==========================
COMPANIES = ["IKIO", "METCO", "SUNSPRINT"]
COMPANY_DB_DIR = "company_db"

load_dotenv()

# --- Ollama (Llama 3) for EVERYTHING (summaries, scoring, justification) ---
OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3:8b-instruct-q4_0")
LLM_TEMPERATURE = float(os.getenv("LLM_TEMPERATURE", "0"))

# Create client (talks to your local ollama daemon)
ollama_client = ollama.Client(host=OLLAMA_HOST)

# ===== Token Budgets =====
LLM_NUM_CTX = int(os.getenv("LLM_NUM_CTX", "8192"))
LLM_INPUT_BUDGET   = int(os.getenv("LLM_INPUT_BUDGET", "6000"))
LLM_OUTPUT_BUDGET  = int(os.getenv("LLM_OUTPUT_BUDGET", "1200"))
EVAL_INPUT_BUDGET  = int(os.getenv("EVAL_INPUT_BUDGET", "6000"))
EVAL_OUTPUT_BUDGET = int(os.getenv("EVAL_OUTPUT_BUDGET", "600"))
MICROCHUNK_INPUT_BUDGET  = int(os.getenv("MICROCHUNK_INPUT_BUDGET", "2500"))
MICROCHUNK_OUTPUT_BUDGET = int(os.getenv("MICROCHUNK_OUTPUT_BUDGET", "500"))
MICROCHUNK_OVERLAP_TOK   = int(os.getenv("MICROCHUNK_OVERLAP_TOK", "150"))

# Retrieval / display knobs
DEFAULT_BATCH_PAGES = 10
TOP_K = 4
SUMMARY_MIN_CHARS = 2500

# ========================== NEW: Company Capabilities & State Licenses ==========================
COMPANY_CAPABILITIES = {
    "IKIO": {
        "scope_capabilities": {
            "lighting": {"supply": True, "supply_installation": True, "percentage": 10},
            "solar": {"supply": True, "supply_installation": True, "percentage": 15},
            "hvac": {"supply": False, "supply_installation": False, "percentage": 0},
            "water_management": {"supply": True, "supply_installation": True, "percentage": 12},
            "emergency_generator": {"supply": True, "supply_installation": True, "percentage": 10},
            "building_envelop": {"supply": False, "supply_installation": False, "percentage": 0},
            "esco": {"supply": True, "supply_installation": True, "percentage": 20}
        },
        "state_licenses": ["CA", "NY", "TX", "FL", "IL", "PA", "OH", "GA", "NC", "MI"],
        "license_percentage": 15
    },
    "METCO": {
        "scope_capabilities": {
            "lighting": {"supply": True, "supply_installation": True, "percentage": 12},
            "solar": {"supply": True, "supply_installation": True, "percentage": 18},
            "hvac": {"supply": True, "supply_installation": True, "percentage": 15},
            "water_management": {"supply": True, "supply_installation": False, "percentage": 8},
            "emergency_generator": {"supply": True, "supply_installation": True, "percentage": 12},
            "building_envelop": {"supply": True, "supply_installation": True, "percentage": 10},
            "esco": {"supply": True, "supply_installation": True, "percentage": 18}
        },
        "state_licenses": ["CA", "TX", "AZ", "NV", "CO", "WA", "OR", "NM", "UT", "ID", "MT", "WY"],
        "license_percentage": 15
    },
    "SUNSPRINT": {
        "scope_capabilities": {
            "lighting": {"supply": True, "supply_installation": False, "percentage": 5},
            "solar": {"supply": True, "supply_installation": True, "percentage": 20},
            "hvac": {"supply": False, "supply_installation": False, "percentage": 0},
            "water_management": {"supply": True, "supply_installation": True, "percentage": 10},
            "emergency_generator": {"supply": True, "supply_installation": False, "percentage": 5},
            "building_envelop": {"supply": False, "supply_installation": False, "percentage": 0},
            "esco": {"supply": True, "supply_installation": True, "percentage": 15}
        },
        "state_licenses": ["FL", "GA", "SC", "NC", "VA", "MD", "DE", "NJ", "AL", "TN"],
        "license_percentage": 15
    }
}

# ========================== Checklist (23) ==========================
CHECKLIST_23 = [
    ("Full solicitation received (RFP, RFQ, drawings, specs, addenda) in the bid document?", "C"),
    ("U.S. federal/state/local procurement requirements reviewed?", "C"),
    ("All bid instructions, deadlines, and addenda understood?", "C"),
    ("Any ambiguities or missing info required? Clarification with the contracting officer/owner required?", "NC"),
    ("What is the project state and is company registered and licensed for work in the project state?", "C"),
    ("Do company meet all mandatory minimum qualifications (experience, capabilities, safety record, etc.) as per the project?", "C"),
    ("Can company meet Diversity/Small Business goals (e.g., SBE, WBE, MBE) if mandatory?", "C"),
    ("Do project require security clearances (for federal/defense projects) and are these obtainable for the company?", "C"),
    ("Is any Environmental/Permitting risks (NEPA, state laws) are required and are these understood and manageable by the company?", "C"),
    ("Is Project scope is clear and achievable with our company's capabilities?", "C"),
    ("Is any Labor relations/Union requirements? If so, are these requirements understood and manageable?", "C"),
    ("Is Site investigation conducted or required or sufficient data (geotech, survey) are available in bid document? If not, what is required?", "NC"),
    ("What are Major technical, commercial, and execution risks required in the bid? Are they identified and manageable by the company?", "C"),
    ("Adequate time and internal resources are available for quality proposal preparation by the company?", "C"),
    ("We have a competitive advantage or unique value proposition for this bid from the company?", "NC"),
    ("Internal team and key partners (especially engineering) are available?", "C"),
    ("Critical vendors, equipment, and subcontractors are available and can comply with project terms?", "C"),
    ("Is any adherence to 'BABA/BAA' or other domestic content rules is required in the BID and is it possible for the company to fulfil this?", "C"),
    ("Is any Contractual Terms & Conditions (indemnity, liability, payment) are required in the BID? Are these acceptable by the company?", "C"),
    ("Is any Required bonds (bid, performance, payment) and insurances in the BID can be secured by the company?", "C"),
    ("Is Preliminary pricing is feasible and within a competitive range for a target profit margin for the company?", "C"),
    ("Is Project aligning with our company's current strategic goals and risk profile?", "C"),
    ("Is Management/Executive buy-in for bidding this project has been secured?", "C"),
]

# ========================== NEW: Scope & License Filtering ==========================
def extract_scope_from_rfp(rfp_summary: dict) -> Dict[str, Any]:
    """
    Extract scope types and work types from RFP summary using keywords.
    Returns: dict with detected scopes and their confidence
    """
    text = (
        rfp_summary.get("summary", "") + " " +
        rfp_summary.get("scope", "") + " " +
        " ".join(rfp_summary.get("key_requirements", [])) + " " +
        " ".join(rfp_summary.get("technical_requirements", []))
    ).lower()
    
    scope_keywords = {
        "lighting": ["lighting", "led", "lamp", "fixture", "illumination", "luminaire"],
        "solar": ["solar", "photovoltaic", "pv", "renewable energy", "solar panel"],
        "hvac": ["hvac", "heating", "ventilation", "air conditioning", "cooling", "climate control"],
        "water_management": ["water", "plumbing", "irrigation", "stormwater", "wastewater"],
        "emergency_generator": ["generator", "emergency power", "backup power", "standby generator"],
        "building_envelop": ["building envelope", "insulation", "weatherization", "roof", "facade"],
        "esco": ["esco", "energy service", "performance contract", "energy savings"]
    }
    
    # Detect work type (supply only vs supply+installation)
    supply_only_keywords = ["supply only", "furnish only", "equipment only", "materials only"]
    installation_keywords = ["install", "installation", "furnish and install", "design-build", "turnkey"]
    
    is_supply_only = any(kw in text for kw in supply_only_keywords)
    has_installation = any(kw in text for kw in installation_keywords)
    
    detected_scopes = {}
    for scope, keywords in scope_keywords.items():
        count = sum(1 for kw in keywords if kw in text)
        if count > 0:
            detected_scopes[scope] = {
                "detected": True,
                "confidence": min(count / len(keywords), 1.0),
                "requires_installation": has_installation and not is_supply_only
            }
    
    return detected_scopes

def extract_state_from_rfp(rfp_summary: dict) -> List[str]:
    """
    Extract state(s) where work will be performed from RFP.
    Returns: List of state abbreviations (e.g., ['CA', 'NY'])
    """
    text = (
        rfp_summary.get("summary", "") + " " +
        rfp_summary.get("scope", "") + " " +
        " ".join(rfp_summary.get("__locations", []))
    )
    
    # US state abbreviations and names
    states_map = {
        "AL": "Alabama", "AK": "Alaska", "AZ": "Arizona", "AR": "Arkansas",
        "CA": "California", "CO": "Colorado", "CT": "Connecticut", "DE": "Delaware",
        "FL": "Florida", "GA": "Georgia", "HI": "Hawaii", "ID": "Idaho",
        "IL": "Illinois", "IN": "Indiana", "IA": "Iowa", "KS": "Kansas",
        "KY": "Kentucky", "LA": "Louisiana", "ME": "Maine", "MD": "Maryland",
        "MA": "Massachusetts", "MI": "Michigan", "MN": "Minnesota", "MS": "Mississippi",
        "MO": "Missouri", "MT": "Montana", "NE": "Nebraska", "NV": "Nevada",
        "NH": "New Hampshire", "NJ": "New Jersey", "NM": "New Mexico", "NY": "New York",
        "NC": "North Carolina", "ND": "North Dakota", "OH": "Ohio", "OK": "Oklahoma",
        "OR": "Oregon", "PA": "Pennsylvania", "RI": "Rhode Island", "SC": "South Carolina",
        "SD": "South Dakota", "TN": "Tennessee", "TX": "Texas", "UT": "Utah",
        "VT": "Vermont", "VA": "Virginia", "WA": "Washington", "WV": "West Virginia",
        "WI": "Wisconsin", "WY": "Wyoming", "DC": "District of Columbia"
    }
    
    detected_states = []
    
    # Search for state abbreviations
    for abbr in states_map.keys():
        if re.search(rf'\b{abbr}\b', text, re.IGNORECASE):
            detected_states.append(abbr)
    
    # Search for full state names
    for abbr, name in states_map.items():
        if name.lower() in text.lower() and abbr not in detected_states:
            detected_states.append(abbr)
    
    return detected_states

def calculate_company_match_percentage(company_name: str, rfp_summary: dict) -> Dict[str, Any]:
    """
    Calculate match percentage based on:
    1. Scope capabilities (Type of work)
    2. State license match
    3. Installation capability if required
    
    Returns: dict with breakdown and total percentage
    """
    if company_name not in COMPANY_CAPABILITIES:
        return {"total_match": 0, "breakdown": {}, "eligible": False, "reason": "Company not in database"}
    
    company_data = COMPANY_CAPABILITIES[company_name]
    detected_scopes = extract_scope_from_rfp(rfp_summary)
    required_states = extract_state_from_rfp(rfp_summary)
    
    # Calculate scope match
    scope_match_total = 0
    scope_breakdown = {}
    eligible = True
    ineligibility_reasons = []
    
    if detected_scopes:
        for scope_name, scope_info in detected_scopes.items():
            capability = company_data["scope_capabilities"].get(scope_name, {})
            
            # Check if company can handle this scope
            if scope_info["requires_installation"]:
                can_do = capability.get("supply_installation", False)
                work_type = "Supply+Installation"
            else:
                can_do = capability.get("supply", False) or capability.get("supply_installation", False)
                work_type = "Supply"
            
            if can_do:
                # Award percentage based on confidence and capability percentage
                awarded_pct = capability.get("percentage", 0) * scope_info["confidence"]
                scope_match_total += awarded_pct
                scope_breakdown[scope_name] = {
                    "required_type": work_type,
                    "can_fulfill": True,
                    "percentage_awarded": round(awarded_pct, 2),
                    "confidence": round(scope_info["confidence"], 2)
                }
            else:
                eligible = False
                ineligibility_reasons.append(f"{scope_name.upper()} ({work_type}) - Company cannot fulfill")
                scope_breakdown[scope_name] = {
                    "required_type": work_type,
                    "can_fulfill": False,
                    "percentage_awarded": 0,
                    "confidence": round(scope_info["confidence"], 2)
                }
    
    # Calculate state license match
    state_match_percentage = 0
    state_breakdown = {}
    
    if required_states:
        licensed_states = company_data.get("state_licenses", [])
        matched_states = [s for s in required_states if s in licensed_states]
        
        if matched_states:
            # Full license percentage if all required states are covered
            state_match_percentage = company_data.get("license_percentage", 15)
            state_breakdown = {
                "required_states": required_states,
                "licensed_states": licensed_states,
                "matched_states": matched_states,
                "percentage_awarded": state_match_percentage,
                "all_states_covered": len(matched_states) == len(required_states)
            }
        else:
            eligible = False
            ineligibility_reasons.append(f"No state license for required states: {', '.join(required_states)}")
            state_breakdown = {
                "required_states": required_states,
                "licensed_states": licensed_states,
                "matched_states": [],
                "percentage_awarded": 0,
                "all_states_covered": False
            }
    else:
        # If no state detected, give benefit of doubt with partial percentage
        state_match_percentage = company_data.get("license_percentage", 15) * 0.5
        state_breakdown = {
            "required_states": ["Not specified in RFP"],
            "licensed_states": company_data.get("state_licenses", []),
            "percentage_awarded": round(state_match_percentage, 2),
            "note": "No specific state requirement detected, partial credit given"
        }
    
    total_match = scope_match_total + state_match_percentage
    
    return {
        "total_match": round(total_match, 2),
        "eligible": eligible,
        "ineligibility_reasons": ineligibility_reasons,
        "breakdown": {
            "scope_match": {
                "total_percentage": round(scope_match_total, 2),
                "details": scope_breakdown
            },
            "state_license_match": state_breakdown
        },
        "detected_scopes": list(detected_scopes.keys()),
        "required_states": required_states
    }

# ========================== Scoring ==========================
def score_item(eval_str: str, criticality: str):
    if eval_str == "Yes":
        return 2 if criticality == "C" else 1
    elif eval_str == "Partial":
        return 1 if criticality == "C" else 0.5
    else:
        return 0

MAX_SCORE = sum(2 if c == "C" else 1 for _, c in CHECKLIST_23)
GO_THRESHOLD = 0.8 * MAX_SCORE

# ========================== Decision Categories ==========================
def get_decision_category(compliance_percent: float) -> str:
    """Return decision category based on compliance percentage."""
    if compliance_percent >= 90:
        return "Strongly-Go"
    elif compliance_percent >= 80:
        return "Go"
    elif compliance_percent >= 70:
        return "Conditional Go"
    elif compliance_percent >= 60:
        return "Evaluate Again"
    else:
        return "No-Go"

def get_decision_color(decision: str) -> str:
    """Return color for decision category."""
    colors = {
        "Strongly-Go": "#28a745",
        "Go": "#5cb85c",
        "Conditional Go": "#ffc107",
        "Evaluate Again": "#ff9800",
        "No-Go": "#dc3545"
    }
    return colors.get(decision, "#6c757d")

def get_decision_emoji(decision: str) -> str:
    """Return emoji for decision category."""
    emojis = {
        "Strongly-Go": "🚀",
        "Go": "✅",
        "Conditional Go": "⚠️",
        "Evaluate Again": "🔍",
        "No-Go": "❌"
    }
    return emojis.get(decision, "❓")

# ========================== OCR (docTR) ==========================
@st.cache_resource
def load_doctr_model(device: str = "cuda"):
    try:
        return ocr_predictor(pretrained=True).to(device)
    except Exception:
        return ocr_predictor(pretrained=True).to("cpu")

def extract_pages_with_doctr(pdf_bytes: bytes, model) -> List[str]:
    """Return a list of page-level texts using docTR OCR."""
    doc = DocumentFile.from_pdf(pdf_bytes)
    result = model(doc)
    return [page.render() for page in result.pages]

def extract_page_texts(file, doctr_model) -> List[str]:
    """
    Try to extract selectable text with PyMuPDF per-page.
    If no text is found (scanned PDF), OCR page-by-page with docTR.
    """
    pages = []
    try:
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                t = page.get_text("text")
                pages.append(t or "")
    except Exception:
        pages = []
    # If most pages are empty => OCR fallback
    if not pages or sum(1 for p in pages if p.strip()) < max(1, int(0.2 * len(pages))):
        file.seek(0)
        pdf_bytes = file.read()
        pages = extract_pages_with_doctr(pdf_bytes, doctr_model)
    file.seek(0)
    return pages

# ========================== Token Utilities ==========================
def _get_token_encoder():
    try:
        import tiktoken
        return tiktoken.get_encoding("cl100k_base")
    except Exception:
        return None

ENC = _get_token_encoder()

def count_tokens(text: str) -> int:
    if not text:
        return 0
    if ENC:
        try:
            return len(ENC.encode(text))
        except Exception:
            pass
    return max(1, int(len(text) / 4))

def trim_text_to_token_limit(text: str, max_tokens: int) -> str:
    """Trim text to fit approx token limit."""
    if count_tokens(text) <= max_tokens:
        return text
    if ENC:
        ids = ENC.encode(text)
        return ENC.decode(ids[:max_tokens])
    approx_chars = max_tokens * 4
    return text[:approx_chars]

def split_text_by_tokens(text: str, max_tokens: int, overlap_tokens: int = 0) -> List[str]:
    """Split text into chunks under max_tokens, with optional overlap."""
    if not text:
        return []
    if ENC:
        ids = ENC.encode(text)
        chunks = []
        i = 0
        step = max_tokens - overlap_tokens if max_tokens > overlap_tokens else max_tokens
        while i < len(ids):
            chunk_ids = ids[i:i + max_tokens]
            chunks.append(ENC.decode(chunk_ids))
            i += step
        return chunks
    approx_chars = max_tokens * 4
    overlap_chars = overlap_tokens * 4
    chunks, i = [], 0
    step = max(1, approx_chars - overlap_chars)
    while i < len(text):
        chunks.append(text[i:i + approx_chars])
        i += step
    return chunks

# ========================== Helpers ==========================
def _dedup_list(lst: List[str]) -> List[str]:
    seen, out = set(), []
    for x in lst:
        if x is None:
            continue
        s = str(x).strip()
        if not s or s in seen:
            continue
        seen.add(s)
        out.append(s)
    return out

def chunk_pages(pages: List[str], batch_pages: int) -> List[List[str]]:
    return [pages[i:i+batch_pages] for i in range(0, len(pages), batch_pages)]

def _compose_long_summary(meta: dict, extracted: dict) -> str:
    pt = meta.get("project_type", "Not specified")
    scope = meta.get("scope", "Not specified")
    domain = meta.get("primary_domain", "Not specified")
    complexity = meta.get("complexity", "Not specified")
    due = meta.get("due_date", "Not specified")
    cost = meta.get("total_cost", "Not specified")

    keys = meta.get("key_requirements", []) or extracted.get("key_requirements", [])
    tech = meta.get("technical_requirements", []) or extracted.get("technical_requirements", [])
    bid = meta.get("bid_requirements", []) or extracted.get("bid_requirements", [])
    kws = meta.get("other_keywords", []) or extracted.get("other_keywords", [])
    dates = extracted.get("dates", [])
    budgets = extracted.get("budget_mentions", [])
    locs = extracted.get("locations", [])

    def bullets(title, arr):
        if not arr: return ""
        lines = "\n".join([f"- {a}" for a in arr])
        return f"\n**{title}:**\n{lines}\n"

    narrative = [
        f"**Project Type:** {pt}",
        f"**Scope:** {scope}",
        f"**Primary Domain:** {domain}",
        f"**Complexity:** {complexity}",
        f"**Due Date:** {due}",
        f"**Estimated Budget / Total Cost:** {cost}",
        "",
        "Detailed consolidated summary assembled from all batch summaries and extracted lists.",
        bullets("Key Requirements", keys),
        bullets("Technical Requirements / Specifications", tech),
        bullets("Bid Instructions & Submission Requirements", bid),
        bullets("Additional Keywords / Entities / Clauses", kws),
        bullets("Dates / Milestones (raw mentions)", dates),
        bullets("Budget Mentions (raw)", budgets),
        bullets("Locations (raw)", locs),
    ]
    return "\n".join([p for p in narrative if p is not None])

# ========================== Ollama JSON helper ==========================
def ollama_json(prompt: str, out_tokens: int) -> dict:
    """Call Ollama (Llama 3) chat API and parse JSON."""
    try:
        resp = ollama_client.chat(
            model=OLLAMA_MODEL,
            messages=[{"role": "user", "content": prompt}],
            options={
                "temperature": LLM_TEMPERATURE,
                "num_ctx": LLM_NUM_CTX,
                "num_predict": int(min(out_tokens, LLM_OUTPUT_BUDGET)),
            },
            stream=False,
        )
        text = resp.get("message", {}).get("content", "")
    except Exception as e:
        return {"error": f"Ollama error: {e}"}

    try:
        return json.loads(text)
    except Exception:
        m = re.search(r'(\{.*\}|\[.*\])', text, flags=re.DOTALL)
        if m:
            try:
                return json.loads(m.group(1))
            except Exception:
                pass
        return {"raw": text.strip()}

# ========================== Micro-summarization (Ollama) ==========================
def micro_summarize_with_llama(long_text: str) -> str:
    """Split into microchunks and produce dense bullet summaries."""
    chunks = split_text_by_tokens(long_text, MICROCHUNK_INPUT_BUDGET, MICROCHUNK_OVERLAP_TOK)
    bullets = []
    for i, ch in enumerate(chunks, start=1):
        ch = trim_text_to_token_limit(ch, MICROCHUNK_INPUT_BUDGET)
        prompt = f"""Summarize the following RFP micro-chunk #{i} into dense bullets only.
Return ONLY JSON: {{"bullets": ["..."]}}

TEXT:
{ch}"""
        data = ollama_json(prompt, MICROCHUNK_OUTPUT_BUDGET)
        lst = data.get("bullets", [])
        if isinstance(lst, list):
            bullets.extend(lst)
        elif isinstance(data, dict) and "raw" in data:
            bullets.append(data["raw"][:500])
    merged = "\n".join(f"- {b}" for b in _dedup_list(bullets))
    return trim_text_to_token_limit(merged, LLM_INPUT_BUDGET)

# ========================== Batch Summaries (Ollama) ==========================
def summarize_batch_with_llama(batch_text: str, batch_idx: int) -> dict:
    """Summarize a ~10-page batch with Llama 3."""
    if count_tokens(batch_text) > LLM_INPUT_BUDGET:
        batch_text = micro_summarize_with_llama(batch_text)
    batch_text = trim_text_to_token_limit(batch_text, LLM_INPUT_BUDGET)

    prompt = f"""You are summarizing a group of RFP pages (batch #{batch_idx}).
Return ONLY JSON with:
{{
 "summary": "dense 400-800 words capturing scope, deliverables, milestones, locations, due dates, budget, bonding/insurance, BABA/BAA or similar, evaluation criteria, submittals, warranty/O&M, risks, unique clauses",
 "key_requirements": ["..."],
 "technical_requirements": ["..."],
 "bid_requirements": ["..."],
 "other_keywords": ["..."],
 "dates": ["..."],
 "budget_mentions": ["..."],
 "locations": ["..."]
}}

TEXT:
{batch_text}
"""
    data = ollama_json(prompt, LLM_OUTPUT_BUDGET)
    for k in ["key_requirements","technical_requirements","bid_requirements","other_keywords","dates","budget_mentions","locations"]:
        data[k] = _dedup_list(data.get(k, []))
    data["summary"] = data.get("summary", "")
    return data

def build_master_summary_with_llama(batch_summaries: List[dict]) -> dict:
    merged = {
        "key_requirements": [],
        "technical_requirements": [],
        "bid_requirements": [],
        "other_keywords": [],
        "dates": [],
        "budget_mentions": [],
        "locations": [],
    }
    parts = []
    for i, b in enumerate(batch_summaries, start=1):
        parts.append(f"=== Batch {i} Summary ===\n{b.get('summary','')}\n")
        for k in merged.keys():
            merged[k].extend(b.get(k, []))

    for k in merged.keys():
        merged[k] = _dedup_list(merged[k])

    batches_text = "\n".join(parts)
    batches_text = trim_text_to_token_limit(batches_text, LLM_INPUT_BUDGET)

    prompt = f"""You are consolidating multiple RFP batch summaries into one master, detailed meta summary.
Return ONLY JSON with these fields:
{{
 "project_type": "...",
 "scope": "...",
 "primary_domain": "...",
 "complexity": "Low|Medium|High",
 "summary": "Comprehensive 3000-5000 characters using information ONLY from the batch summaries and lists. Integrate scope, deliverables, milestones, sites/locations, due dates, budget ranges, bonding/insurance, BABA/BAA or similar, evaluation criteria, submittals, warranty/O&M, risks, and unique clauses. Deduplicate and be specific.",
 "due_date": "...",
 "total_cost": "...",
 "key_requirements": [],
 "technical_requirements": [],
 "bid_requirements": [],
 "other_keywords": []
}}

Merged lists to integrate (deduplicated):
key_requirements: {merged["key_requirements"]}
technical_requirements: {merged["technical_requirements"]}
bid_requirements: {merged["bid_requirements"]}
other_keywords: {merged["other_keywords"]}
dates: {merged["dates"]}
budget_mentions: {merged["budget_mentions"]}
locations: {merged["locations"]}

BATCH SUMMARIES (for reference, do not repeat verbatim):
{batches_text}
"""
    meta = ollama_json(prompt, LLM_OUTPUT_BUDGET)

    for list_key in ["key_requirements","technical_requirements","bid_requirements","other_keywords"]:
        meta[list_key] = _dedup_list((meta.get(list_key) or []) + merged[list_key])

    meta["__dates"] = merged["dates"]
    meta["__budget_mentions"] = merged["budget_mentions"]
    meta["__locations"] = merged["locations"]

    if len(meta.get("summary","")) < SUMMARY_MIN_CHARS:
        meta = {**meta}
        meta["summary"] = _compose_long_summary(meta, {
            "key_requirements": merged["key_requirements"],
            "technical_requirements": merged["technical_requirements"],
            "bid_requirements": merged["bid_requirements"],
            "other_keywords": merged["other_keywords"],
            "dates": merged["dates"],
            "budget_mentions": merged["budget_mentions"],
            "locations": merged["locations"],
        })
    return meta

# ========================== DOCX Export ==========================
def export_to_docx(company_results: dict, best_company: str, file_path: str, rfp_summary: dict, justification: str | None, match_percentages: dict):
    doc = Document()
    doc.add_heading("Bid Compliance Evaluation Report (23-Point Weighted Scoring)", level=1)

    # RFP Summary
    doc.add_heading("RFP Summary", level=2)
    doc.add_paragraph(f"Project Type: {rfp_summary.get('project_type','N/A')}")
    doc.add_paragraph(f"Scope: {rfp_summary.get('scope','N/A')}")
    doc.add_paragraph(f"Primary Domain: {rfp_summary.get('primary_domain','N/A')}")
    doc.add_paragraph(f"Complexity: {rfp_summary.get('complexity','N/A')}")
    doc.add_paragraph(f"Due Date: {rfp_summary.get('due_date','N/A')}")
    doc.add_paragraph(f"Estimated Cost: {rfp_summary.get('total_cost','N/A')}")
    
    # Add scope and license analysis
    doc.add_heading("Scope & License Analysis", level=2)
    for name, match_data in match_percentages.items():
        doc.add_heading(f"{name} - Eligibility Analysis", level=3)
        doc.add_paragraph(f"Total Match Score: {match_data['total_match']}%")
        doc.add_paragraph(f"Eligible: {'✅ Yes' if match_data['eligible'] else '❌ No'}")
        if match_data.get('ineligibility_reasons'):
            doc.add_paragraph(f"Reasons for Ineligibility: {', '.join(match_data['ineligibility_reasons'])}")
        
        # Scope breakdown
        scope_details = match_data['breakdown']['scope_match']['details']
        if scope_details:
            doc.add_paragraph("Scope Capabilities:")
            for scope_name, scope_info in scope_details.items():
                doc.add_paragraph(f"  • {scope_name.upper()}: {'✓' if scope_info['can_fulfill'] else '✗'} - {scope_info['percentage_awarded']}%")
        
        # State license breakdown
        state_info = match_data['breakdown']['state_license_match']
        doc.add_paragraph(f"State License Match: {state_info.get('percentage_awarded', 0)}%")
        doc.add_paragraph(f"Required States: {', '.join(state_info.get('required_states', ['N/A']))}")
        doc.add_paragraph(f"Matched States: {', '.join(state_info.get('matched_states', ['None']))}")
    
    doc.add_paragraph("Executive Summary:")
    doc.add_paragraph(rfp_summary.get('summary','N/A'))

    for name, evaluation in company_results.items():
        doc.add_heading(f"{name} – Detailed Evaluation", level=2)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text, hdr[5].text = (
            "Requirement", "Criticality", "Evaluation", "Rationale", "Score", "Context Used"
        )
        for item in evaluation:
            row = table.add_row().cells
            row[0].text = item["requirement"]
            row[1].text = item["criticality"]
            row[2].text = item["evaluation"]
            row[3].text = item["comments"]
            row[4].text = str(item["score"])
            row[5].text = item.get("context_source","")
        fs = sum(x["score"] for x in evaluation)
        compliance = round((fs / MAX_SCORE) * 100, 2)
        decision = get_decision_category(compliance)
        emoji = get_decision_emoji(decision)
        doc.add_paragraph(f"Total Score: {fs} / {MAX_SCORE}")
        doc.add_paragraph(f"Compliance %: {compliance}%")
        doc.add_paragraph(f"Decision: {emoji} {decision}")
        doc.add_page_break()

    doc.add_heading("🏆 Best Company Recommendation", level=2)
    doc.add_paragraph(f"The best company for this RFP is: **{best_company}**")
    if justification:
        doc.add_heading("Executive Justification", level=2)
        doc.add_paragraph(justification)
    doc.save(file_path)

# ========================== Scorer (Llama 3) ==========================
class WeightedBidScorer23:
    def __init__(self, checklist, emb_model):
        self.checklist = checklist
        self.emb_model = emb_model

    def retrieve_relevant_context(self, company_texts: List[str], requirement: str) -> str:
        if not company_texts:
            return ""
        req_vec = self.emb_model.encode([requirement], convert_to_numpy=True)
        text_vecs = self.emb_model.encode(company_texts, convert_to_numpy=True)

        req_norm = np.linalg.norm(req_vec, axis=1, keepdims=True) + 1e-12
        txt_norms = np.linalg.norm(text_vecs, axis=1, keepdims=True) + 1e-12
        sims = (text_vecs @ req_vec.T)[:, 0] / (txt_norms[:, 0] * req_norm[0, 0])

        top_idx = sims.argsort()[-TOP_K:][::-1]
        snippets = []
        per_snippet_tok = max(200, int((EVAL_INPUT_BUDGET // max(1, TOP_K)) * 0.4))
        for i in top_idx:
            s = trim_text_to_token_limit(company_texts[i], per_snippet_tok)
            snippets.append(s)
        return "\n\n--- COMPANY CHUNK ---\n\n".join(snippets)

    def _assemble_rfp_context(self, rfp_summary: dict) -> str:
        def cap_list(lst, n=50):
            return lst[:n] if isinstance(lst, list) else []

        ctx = (
            "RFP Summary (long):\n" + rfp_summary.get("summary","") + "\n\n"
            f"Key requirements: {cap_list(rfp_summary.get('key_requirements',[]))}\n"
            f"Technical requirements: {cap_list(rfp_summary.get('technical_requirements',[]))}\n"
            f"Bid requirements: {cap_list(rfp_summary.get('bid_requirements',[]))}\n"
            f"Other keywords: {cap_list(rfp_summary.get('other_keywords',[]))}\n"
            f"Dates: {cap_list(rfp_summary.get('__dates',[]), 40)} | "
            f"Budget mentions: {cap_list(rfp_summary.get('__budget_mentions',[]), 40)} | "
            f"Locations: {cap_list(rfp_summary.get('__locations',[]), 40)}"
        )
        return trim_text_to_token_limit(ctx, EVAL_INPUT_BUDGET)

    def evaluate_item(self, requirement: str, criticality: str, rfp_summary: dict,
                      company_name: str, company_texts: List[str], check_rfp_only=False) -> dict:
        if check_rfp_only:
            context_source = "RFP-only"
            context = self._assemble_rfp_context(rfp_summary)
        else:
            context_source = "RFP+Company"
            rfp_ctx = self._assemble_rfp_context(rfp_summary)
            company_ctx = self.retrieve_relevant_context(company_texts, requirement)
            merged = rfp_ctx + "\n\nCompany Evidence:\n" + company_ctx
            context = trim_text_to_token_limit(merged, EVAL_INPUT_BUDGET)

        prompt = f"""
Evaluate requirement for company "{company_name}".

Requirement: {requirement} ({criticality})
Context:
{context}

Rules:
- Use ONLY the provided context.
- If evidence is insufficient, evaluation="Unknown", comments="Unknown (not in documents)".
- Return ONLY JSON:
{{
 "requirement":"{requirement}",
 "criticality":"{criticality}",
 "evaluation":"Yes|Partial|No|Unknown",
 "comments":"short rationale tied to the evidence"
}}
"""
        data = ollama_json(prompt, EVAL_OUTPUT_BUDGET)
        requirement_val = data.get("requirement", requirement)
        criticality_val = data.get("criticality", criticality)
        evaluation_val = data.get("evaluation", "Unknown")
        comments_val = data.get("comments", "Unknown (not in documents)")

        result = {
            "requirement": requirement_val,
            "criticality": criticality_val,
            "evaluation": evaluation_val,
            "comments": comments_val,
            "score": score_item(evaluation_val, criticality_val),
            "context_source": context_source
        }
        return result

    def score_company(self, rfp_summary: dict, company_name: str, company_texts: List[str]) -> List[Dict[str, Any]]:
        evaluations = []
        for idx, (req, crit) in enumerate(self.checklist):
            check_rfp_only = idx < 4
            evaluations.append(
                self.evaluate_item(req, crit, rfp_summary, company_name, company_texts, check_rfp_only)
            )
        return evaluations

# ========================== Final Justification (Llama 3) ==========================
def generate_final_justification(best_company, results, rfp_summary, match_percentages):
    best_match = match_percentages.get(best_company, {})
    
    facts = {
        "best_company": best_company,
        "max_score": MAX_SCORE,
        "rfp_due": rfp_summary.get("due_date","Not specified"),
        "rfp_scope": rfp_summary.get("scope","Not specified"),
        "top_requirements": rfp_summary.get("key_requirements", [])[:12],
        "scope_match_percentage": best_match.get("total_match", 0),
        "eligible": best_match.get("eligible", False),
        "detected_scopes": best_match.get("detected_scopes", []),
        "required_states": best_match.get("required_states", []),
        "scores": {
            c: sum(it["score"] for it in evals)
            for c, evals in results.items()
        },
        "decisive_items": {
            c: sorted(
                [(it["requirement"], it["evaluation"], it["score"]) for it in evals],
                key=lambda x: x[2], reverse=True
            )[:6]
            for c, evals in results.items()
        }
    }

    prompt = f"""
Write an executive justification (6–10 bullets) explaining why "{best_company}" is recommended.
Use only the facts below. Emphasize decisive checklist items, scope/license match percentage, and any material risks/gaps.
Be board-ready, specific, and concise.

FACTS:
{json.dumps(facts, ensure_ascii=False)}
"""
    data = ollama_json(prompt, 500)
    if isinstance(data, dict):
        if "bullets" in data:
            return "\n".join(f"- {b}" for b in data["bullets"])
        if "text" in data:
            return str(data["text"])
        if "raw" in data:
            return str(data["raw"])
    return str(data)

# ========================== Init ==========================
device = "cuda" if torch.cuda.is_available() else "cpu"
gpu_available = torch.cuda.is_available()
gpu_name = torch.cuda.get_device_name(0) if gpu_available else "N/A"
gpu_memory = f"{torch.cuda.get_device_properties(0).total_memory / 1e9:.2f} GB" if gpu_available else "N/A"

emb_model = SentenceTransformer("all-MiniLM-L6-v2", device=device)
doctr_model = load_doctr_model("cuda" if gpu_available else "cpu")
scorer23 = WeightedBidScorer23(CHECKLIST_23, emb_model)

# ========================== Streamlit UI ==========================
st.title("🎯 Advanced Bid Analyzer with Scope & License Filtering")
st.caption("10-Page Batching | Ollama Llama3 | GPU-Accelerated | Intelligent Filtering")

# GPU Status
col1, col2, col3 = st.columns(3)
with col1:
    if gpu_available:
        st.success(f"🎮 GPU: {gpu_name}")
    else:
        st.warning("⚠️ GPU: Not Available (Using CPU)")
with col2:
    st.info(f"💾 VRAM: {gpu_memory}")
with col3:
    st.info(f"📊 Device: {device.upper()}")

st.info("""
**📊 Filtering Criteria:** 
1. **Type of Work/Scope:** Lighting, Solar, HVAC, Water Management, Emergency Generator, Building Envelope, ESCO
2. **State Licenses:** Match required project states with company licenses
3. **Work Type:** Supply only vs. Supply+Installation capability
""")

# Show company capabilities
with st.expander("🏢 View Company Capabilities & State Licenses"):
    for company, data in COMPANY_CAPABILITIES.items():
        st.markdown(f"### {company}")
        st.markdown(f"**Licensed States:** {', '.join(data['state_licenses'])}")
        st.markdown("**Scope Capabilities:**")
        for scope, cap in data["scope_capabilities"].items():
            supply = "✓" if cap["supply"] else "✗"
            install = "✓" if cap["supply_installation"] else "✗"
            pct = cap["percentage"]
            st.markdown(f"- {scope.upper()}: Supply {supply} | Installation {install} | Match %: {pct}%")
        st.markdown("---")

rfp_file = st.file_uploader("Upload RFP PDF", type="pdf")

if rfp_file and st.button("🚀 Analyze RFP with Filtering"):
    if gpu_available:
        torch.cuda.empty_cache()
    
    # Extract pages
    with st.spinner("Extracting per-page text (OCR if needed)..."):
        pages = extract_page_texts(rfp_file, doctr_model)
    st.info(f"Total pages detected: {len(pages)}")
    
    if gpu_available:
        torch.cuda.empty_cache()

    # Summarize in batches
    batches = chunk_pages(pages, int(DEFAULT_BATCH_PAGES))
    batch_summaries = []
    with st.spinner(f"Summarizing in {len(batches)} batch(es)..."):
        for i, grp in enumerate(batches, start=1):
            batch_text = "\n\n--- PAGE BREAK ---\n\n".join(grp)
            data = summarize_batch_with_llama(batch_text, i)
            batch_summaries.append(data)

    # Build master summary
    with st.spinner("Building master RFP summary..."):
        rfp_summary = build_master_summary_with_llama(batch_summaries)

    # ===== NEW: Scope & License Filtering =====
    st.markdown("---")
    st.subheader("🔍 Scope & License Match Analysis")
    
    match_percentages = {}
    eligible_companies = []
    
    for company in COMPANIES:
        match_data = calculate_company_match_percentage(company, rfp_summary)
        match_percentages[company] = match_data
        
        if match_data["eligible"]:
            eligible_companies.append(company)
        
        # Display match results
        if match_data["eligible"]:
            st.success(f"✅ **{company}** - Eligible | Match Score: **{match_data['total_match']}%**")
        else:
            st.error(f"❌ **{company}** - Not Eligible | Reasons: {', '.join(match_data['ineligibility_reasons'])}")
        
        with st.expander(f"📊 {company} - Detailed Breakdown"):
            st.json(match_data["breakdown"])
    
    if not eligible_companies:
        st.warning("⚠️ No companies are eligible for this RFP based on scope and license requirements.")
        st.stop()
    
    st.info(f"✅ **Eligible Companies:** {', '.join(eligible_companies)}")
    
    # Present RFP Summary
    st.markdown("---")
    st.subheader("📑 RFP Summary")
    st.markdown(f"- **Project Type:** {rfp_summary.get('project_type','Not specified')}")
    st.markdown(f"- **Scope:** {rfp_summary.get('scope','Not specified')}")
    st.markdown(f"- **Primary Domain:** {rfp_summary.get('primary_domain','Not specified')}")
    st.markdown(f"- **Complexity:** {rfp_summary.get('complexity','Not specified')}")
    st.markdown(f"- **Due Date:** {rfp_summary.get('due_date','Not specified')}")
    st.markdown(f"- **Estimated Cost:** {rfp_summary.get('total_cost','Not specified')}")
    
    # Detected scopes and states
    detected_scopes = extract_scope_from_rfp(rfp_summary)
    required_states = extract_state_from_rfp(rfp_summary)
    st.markdown(f"- **Detected Scopes:** {', '.join([s.upper() for s in detected_scopes.keys()]) if detected_scopes else 'None detected'}")
    st.markdown(f"- **Required States:** {', '.join(required_states) if required_states else 'Not specified'}")

    # Evaluation (only for eligible companies)
    results, totals = {}, {}
    for company in eligible_companies:
        path = os.path.join(COMPANY_DB_DIR, f"{company}.json")
        if not os.path.exists(path):
            st.warning(f"⚠️ {company} knowledge base missing at {path}.")
            continue
        with open(path, "r", encoding="utf-8") as f:
            db = json.load(f)
        company_texts = db.get("texts", [])

        with st.spinner(f"Scoring {company}..."):
            evaluations = scorer23.score_company(rfp_summary, company, company_texts)
            results[company] = evaluations
            
            # Add match percentage to total score
            checklist_score = sum(x["score"] for x in evaluations)
            match_bonus = match_percentages[company]["total_match"]
            
            # Store both scores
            totals[company] = {
                "checklist_score": checklist_score,
                "match_percentage": match_bonus,
                "total_score": checklist_score + match_bonus
            }
            
            if gpu_available:
                torch.cuda.empty_cache()

    if results:
        st.markdown("---")
        st.subheader("📊 Company Comparison Summary")
        
        # Quick visual summary cards
        cols = st.columns(len(results))
        for idx, (company, evals) in enumerate(results.items()):
            checklist_score = totals[company]["checklist_score"]
            match_bonus = totals[company]["match_percentage"]
            total_score = totals[company]["total_score"]
            
            # Calculate compliance based on checklist only
            compliance = round((checklist_score / MAX_SCORE) * 100, 2)
            decision = get_decision_category(compliance)
            emoji = get_decision_emoji(decision)
            
            with cols[idx]:
                st.markdown(f"""
                <div style="padding: 15px; border-radius: 10px; background-color: {get_decision_color(decision)}20; border: 2px solid {get_decision_color(decision)};">
                    <h3 style="text-align: center; margin: 0;">{company}</h3>
                    <p style="text-align: center; font-size: 2em; margin: 10px 0;">{emoji}</p>
                    <p style="text-align: center; font-weight: bold; margin: 5px 0;">{compliance}%</p>
                    <p style="text-align: center; color: {get_decision_color(decision)}; font-weight: bold; margin: 5px 0;">{decision}</p>
                    <p style="text-align: center; font-size: 0.9em; margin: 5px 0;">Checklist: {checklist_score}/{MAX_SCORE}</p>
                    <p style="text-align: center; font-size: 0.9em; margin: 5px 0;">Match Bonus: +{match_bonus}%</p>
                    <p style="text-align: center; font-size: 1.1em; font-weight: bold; margin: 5px 0;">Total: {total_score:.2f}</p>
                </div>
                """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Build comparison data
        comp_data = []
        for c, evals in results.items():
            score_data = totals[c]
            compliance = round((score_data["checklist_score"] / MAX_SCORE) * 100, 2)
            decision = get_decision_category(compliance)
            emoji = get_decision_emoji(decision)
            comp_data.append({
                "Company": c,
                "Checklist Score": f"{score_data['checklist_score']}/{MAX_SCORE}",
                "Match %": f"+{score_data['match_percentage']}",
                "Total Score": round(score_data["total_score"], 2),
                "Compliance %": compliance,
                "Decision": f"{emoji} {decision}"
            })
        
        comp_df = pd.DataFrame(comp_data).sort_values("Total Score", ascending=False)
        st.dataframe(comp_df, use_container_width=True, hide_index=True)
        
        # Best company based on total score
        best = max(totals, key=lambda c: totals[c]["total_score"])
        best_data = totals[best]
        best_compliance = round((best_data["checklist_score"] / MAX_SCORE) * 100, 2)
        best_decision = get_decision_category(best_compliance)
        best_emoji = get_decision_emoji(best_decision)
        
        st.markdown("---")
        st.success(f"🏆 **Best Company: {best}** | {best_emoji} {best_decision} ({best_compliance}%) | Total Score: {best_data['total_score']:.2f}")
        
        # Detailed scoring
        st.markdown("---")
        with st.expander("📂 View Detailed Scoring for Each Company"):
            for company, evals in results.items():
                score_data = totals[company]
                compliance = round((score_data["checklist_score"] / MAX_SCORE) * 100, 2)
                decision = get_decision_category(compliance)
                emoji = get_decision_emoji(decision)
                
                st.markdown(f"### {company} - {emoji} {decision} ({compliance}%)")
                st.markdown(f"**Checklist Score:** {score_data['checklist_score']}/{MAX_SCORE}")
                st.markdown(f"**Scope/License Match Bonus:** +{score_data['match_percentage']}%")
                st.markdown(f"**Total Score:** {score_data['total_score']:.2f}")
                
                df = pd.DataFrame(evals)
                st.dataframe(df[["requirement","criticality","evaluation","comments","score","context_source"]],
                             use_container_width=True)
                st.markdown("---")

        # Final justification
        with st.spinner("Generating executive justification..."):
            justification = generate_final_justification(best, results, rfp_summary, match_percentages)

        if justification:
            st.markdown("### 🧭 Executive Justification")
            st.write(justification)

        # Export DOCX
        output_file = "Bid_Report_Enhanced_Filtering.docx"
        export_to_docx(results, best, output_file, rfp_summary, justification, match_percentages)
        with open(output_file, "rb") as f:
            st.download_button("⬇️ Download Full Report", f, file_name=output_file)

